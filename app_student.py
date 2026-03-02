import streamlit as st
import pandas as pd
import plotly.graph_objects as go
import plotly.express as px
from streamlit_option_menu import option_menu
import openai
import os
import functools
import io
import re

# ==============================================================================
# 1. 页面基础配置 (SaaS 宽屏模式)
# ==============================================================================
st.set_page_config(page_title="英华学情查询系统", layout="wide", page_icon="🎓", initial_sidebar_state="expanded")

# ==============================================================================
# 🎨 顶级 SaaS 美学 CSS 样式注入 (极简、紧凑、去白边)
# ==============================================================================
st.markdown("""
<style>
    #MainMenu {visibility: hidden;} 
    footer {visibility: hidden;} 
    header {background: transparent !important;}
    
    .block-container { padding-top: 1rem !important; padding-bottom: 2rem !important; max-width: 96% !important;}
    .stApp { background-color: #F4F7F9; }
    
    /* =========================================================
       1. 侧边栏瘦身与底色
       ========================================================= */
    [data-testid="stSidebar"] {
        min-width: 250px !important; 
        max-width: 250px !important;
        background-color: #0B1120 !important;
    }
    [data-testid="stSidebar"] p, [data-testid="stSidebar"] label, [data-testid="stSidebar"] span {
        color: #94A3B8 !important; 
    }
    [data-testid="stSidebar"] h1, [data-testid="stSidebar"] h2, [data-testid="stSidebar"] h3 {
        color: #F8FAFC !important; 
    }
    [data-testid="stSidebar"] hr {
        border-color: rgba(255,255,255,0.05) !important;
        margin-top: 10px !important; margin-bottom: 10px !important;
    }

    /* =========================================================
       2. 下拉框 (Selectbox) 彻底暗黑化
       ========================================================= */
    [data-testid="stSidebar"] div[data-testid="stSelectbox"] {
        margin-bottom: -15px !important; 
    }
    [data-testid="stSidebar"] div[data-baseweb="select"] > div {
        background-color: #1E293B !important; 
        border: 1px solid rgba(255,255,255,0.05) !important;
        border-radius: 8px !important;
    }
    [data-testid="stSidebar"] div[data-baseweb="select"] div {
        color: #E2E8F0 !important; 
    }
    [data-testid="stSidebar"] div[data-baseweb="select"] svg {
        fill: #94A3B8 !important;
    }
    ul[role="listbox"] {
        background-color: #1E293B !important;
    }
    li[role="option"] {
        background-color: #1E293B !important;
        color: #E2E8F0 !important;
    }
    li[role="option"]:hover {
        background-color: #334155 !important;
    }

    /* =========================================================
       3. 导航折叠面板：去白底、去毛边、极致紧凑
       ========================================================= */
    [data-testid="stSidebar"] [data-testid="stExpander"] {
        border: none !important;
        background: transparent !important;
    }
    [data-testid="stSidebar"] [data-testid="stExpander"] details, 
    [data-testid="stSidebar"] [data-testid="stExpander"] summary {
        background: transparent !important;
        border: none !important;
        padding: 0 !important;
        margin-bottom: 5px !important;
    }
    [data-testid="stSidebar"] [data-testid="stExpander"] summary p {
        font-size: 15px !important;
        font-weight: bold !important;
        color: #94A3B8 !important;
    }
    [data-testid="stSidebar"] [data-testid="stExpander"] summary svg {
        fill: #64748B !important;
    }
    [data-testid="stSidebar"] [data-testid="stExpander"] [data-testid="stExpanderDetails"] {
        background-color: #162032 !important; 
        border-radius: 10px !important;
        padding: 0 !important; 
        overflow: hidden !important; 
        border: 1px solid rgba(255,255,255,0.03) !important;
    }
    [data-testid="stSidebar"] iframe {
        background-color: transparent !important;
    }

    /* =========================================================
       4. 退出按钮暗黑化
       ========================================================= */
    [data-testid="stSidebar"] div.stButton > button {
        background-color: #1E293B !important; 
        color: #94A3B8 !important; 
        border: 1px solid rgba(255,255,255,0.05) !important;
        border-radius: 8px !important;
        font-weight: bold !important;
        padding-top: 10px !important;
        padding-bottom: 10px !important;
        transition: all 0.3s ease;
    }
    [data-testid="stSidebar"] div.stButton > button:hover {
        background-color: #334155 !important; 
        color: #F8FAFC !important; 
        border-color: rgba(255,255,255,0.1) !important;
    }

    /* =========================================================
       5. 右侧内容区美化
       ========================================================= */
    .header-card {
        background-color: #FFFFFF; padding: 15px 25px; border-radius: 10px;
        box-shadow: 0 2px 10px rgba(0,0,0,0.03); display: flex; justify-content: space-between;
        align-items: center; margin-bottom: 20px; border-left: 6px solid #3B82F6;
    }
    div[data-testid="stMetric"] {
        background-color: #FFFFFF; border-radius: 10px; padding: 15px;
        box-shadow: 0 2px 10px rgba(0,0,0,0.02); border: 1px solid #F1F5F9; text-align: center;
    }
    div[data-testid="stMetric"] label { font-size: 14px !important; color: #64748B !important; }
    div[data-testid="stMetric"] div[data-testid="stMetricValue"] { font-size: 26px !important; color: #0F172A !important; font-weight: 800 !important; }
    div[data-testid="stForm"] { background-color: #ffffff; padding: 40px; border-radius: 15px; box-shadow: 0 8px 20px rgba(0,0,0,0.05); border: none; }
    .congrats-banner { background: linear-gradient(90deg, #FFFBEB, #FFF7ED); border: 2px solid #FCD34D; color: #92400E; padding: 12px 20px; border-radius: 12px; text-align: center; font-size: 18px; font-weight: bold; margin-bottom: 25px; box-shadow: 0 4px 12px rgba(252, 211, 77, 0.2); line-height: 1.6; }
    .main-title { text-align: center; color: #0F172A; font-size: 28px; font-weight: 800; margin-bottom: 15px; }
    .ai-box { background: #FFFFFF; border-left: 5px solid #3B82F6; padding: 20px; border-radius: 8px; font-size: 15px; color: #333; line-height: 1.8; box-shadow: 0 2px 10px rgba(0,0,0,0.03);}
</style>
""", unsafe_allow_html=True)

CHART_CONFIG = {'displayModeBar': False, 'scrollZoom': False}

if 'current_grade' not in st.session_state: st.session_state.current_grade = "高三"
if 'logged_in_student' not in st.session_state: st.session_state.logged_in_student = None
if 'logged_in_id' not in st.session_state: st.session_state.logged_in_id = None
if 'logged_in_direction' not in st.session_state: st.session_state.logged_in_direction = None

def logout():
    for key in ['logged_in_student', 'logged_in_id', 'logged_in_direction']: st.session_state[key] = None
    st.rerun()

# ==============================================================================
# 👑 核心引擎：动态读取配置与数据
# ==============================================================================
try:
    AI_API_KEY = st.secrets.get("DEEPSEEK_API_KEY", "")
    URL_EXAM_CONFIG = st.secrets.get("URL_EXAM_CONFIG", "")
except Exception as e:
    st.error("⚠️ 系统配置读取失败，请检查后台 Secrets。")
    st.stop()

if AI_API_KEY: client = openai.OpenAI(api_key=AI_API_KEY, base_url="https://api.deepseek.com")
else: client = None

def clean_url(url):
    if pd.isna(url): return ""
    u = str(url).strip()
    if u.lower() == 'nan': return ""
    return u

@st.cache_data(ttl=300) 
def load_exam_config(url):
    try: return pd.read_csv(url, on_bad_lines='skip')
    except: return pd.DataFrame()

def normalize_class_name(c):
    if pd.isna(c): return ""
    c = str(c).replace(" ", "").strip()
    mapping = {'1':'一','2':'二','3':'三','4':'四','5':'五','6':'六','7':'七','8':'八','9':'九','0':'零'}
    for k, v in mapping.items(): c = c.replace(k, v)
    c = c.replace("高三","").replace("高二","").replace("高一","").replace("年级","").replace("()","").replace("（）","")
    if not c.endswith("班"): c += "班"
    return c

def clean_str(val):
    if pd.isna(val): return ""
    v = str(val).strip()
    if v.endswith('.0'): v = v[:-2]
    return v

def clean_name(val):
    if pd.isna(val): return ""
    return str(val).replace(" ", "").strip()

@st.cache_data(ttl=600)
def load_data(url, header_lines=0):
    if not url or not url.strip(): return None
    try: return pd.read_csv(url, header=header_lines, on_bad_lines='skip')
    except: return None

@st.cache_data(ttl=600, show_spinner=False)
def build_master_df(grade_key):
    config_df = load_exam_config(URL_EXAM_CONFIG)
    exams = []
    if not config_df.empty and '年级' in config_df.columns:
        grade_config = config_df[config_df['年级'].astype(str).str.strip() == grade_key]
        for _, row in grade_config.iterrows():
            exams.append({
                "name": str(row.get('考试名称', '')).strip(),
                "语文": clean_url(row.get('语文')), "数学": clean_url(row.get('数学')),
                "英语": clean_url(row.get('英语')), "物理": clean_url(row.get('物理')),
                "化学": clean_url(row.get('化学')), "生物": clean_url(row.get('生物')),
                "历史": clean_url(row.get('历史')), "政治": clean_url(row.get('政治')),
                "地理": clean_url(row.get('地理'))
            })
    if not exams: return None, None, []
    
    latest_exam = exams[-1]
    dfs = []
    subs = ['语文','数学','英语','物理','化学','生物','历史','政治','地理']
    for sub in subs:
        url = latest_exam.get(sub)
        if url:
            df_sub = load_data(url, header_lines=[0,1,2])
            if df_sub is not None:
                name_c, id_c, cls_c = None, None, None
                for col in df_sub.columns:
                    cstr = str(col[0]) if isinstance(col, tuple) else str(col)
                    if '姓名' in cstr: name_c = col
                    elif '考号' in cstr or '学号' in cstr: id_c = col
                    elif '班级' in cstr: cls_c = col
                if name_c and id_c:
                    res = []
                    for _, row in df_sub.iterrows():
                        tot = 0
                        for c in df_sub.columns:
                            if c in [name_c, id_c, cls_c]: continue
                            cstr = str(c[0]) if isinstance(c, tuple) else str(c)
                            if '总分' in cstr or '排名' in cstr: continue
                            try: 
                                val = float(row[c])
                                if pd.notna(val): tot += val
                            except: pass
                        s_name = clean_name(row[name_c])
                        s_id = clean_str(row[id_c])
                        s_cls = normalize_class_name(row[cls_c]) if cls_c else "未分班"
                        if s_id: res.append({'姓名': s_name, '考号': s_id, '班级': s_cls, sub: round(tot, 1)})
                    if res: dfs.append(pd.DataFrame(res))
    if not dfs: return None, latest_exam, exams
    master = functools.reduce(lambda l, r: pd.merge(l, r, on=['姓名','考号','班级'], how='outer'), dfs)
    present_subs = [s for s in subs if s in master.columns]
    master[present_subs] = master[present_subs].fillna(0)
    master['总分'] = master[present_subs].sum(axis=1).round(1)
    
    def get_dir(r):
        if r.get('物理', 0) > 0: return "物理方向"
        if r.get('历史', 0) > 0: return "历史方向"
        cls_n = str(r.get('班级', ''))
        if '文' in cls_n or '史' in cls_n: return "历史方向"
        if '理' in cls_n or '物' in cls_n: return "物理方向"
        return "综合方向"
    master['方向'] = master.apply(get_dir, axis=1)
    master['总分班级排名'] = master.groupby(['班级', '方向'])['总分'].rank(ascending=False, method='min').fillna(0).astype(int)
    master['总分年级排名'] = master.groupby(['方向'])['总分'].rank(ascending=False, method='min').fillna(0).astype(int)
    return master, latest_exam, exams

# ==============================================================================
# 🎨 导出与排版模块 (压缩表格)
# ==============================================================================
def render_html_table(df):
    html = """
    <div style="width: 100%; overflow-x: auto; margin-bottom: 25px; border-radius: 8px; box-shadow: 0 2px 8px rgba(0,0,0,0.03); border: 1px solid #E8E8E8; background: white;">
    <table style="width: 100%; border-collapse: collapse; font-size: 14px; text-align: center; font-family: 'Helvetica Neue', Arial, sans-serif;">
    """
    html += "<tr>" + "".join([f"<th style='background-color: #FAFAFA; color: #333; padding: 10px 12px; border-bottom: 2px solid #E8E8E8; white-space: nowrap; font-weight: 800;'>{col}</th>" for col in df.columns]) + "</tr>"
    for i, row in df.iterrows():
        bg_color = "#FFFFFF" if i % 2 == 0 else "#FAFAFA"
        html += f"<tr style='background-color: {bg_color}; transition: background-color 0.2s;' onmouseover=\"this.style.backgroundColor='#E6F7FF'\" onmouseout=\"this.style.backgroundColor='{bg_color}'\">"
        for col in df.columns:
            val = row[col]
            if isinstance(val, float): val = f"{val:.1f}"
            html += f"<td style='padding: 8px 12px; border-bottom: 1px solid #F0F0F0; color: #555;'>{val}</td>"
        html += "</tr>"
    html += "</table></div>"
    st.markdown(html, unsafe_allow_html=True)

def generate_ai_doc(title, content):
    try:
        import docx
        from docx.shared import Pt
        from docx.oxml.ns import qn
        from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
        doc = docx.Document()
        style = doc.styles['Normal']
        style.font.name = '仿宋'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        style.font.size = Pt(10.5)
        
        h = doc.add_heading(level=1)
        h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = h.add_run(title)
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(14)
        run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)

        def add_runs_to_paragraph(p, text):
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    clean_text = part[2:-2].replace('*', '') 
                    r = p.add_run(clean_text)
                    r.bold = True
                else:
                    clean_text = part.replace('*', '')
                    r = p.add_run(clean_text)
                r.font.name = '仿宋'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
                r.font.size = Pt(10.5)

        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line: continue
            if line.startswith('- '): line = line[2:].strip()
            if line.startswith('* '): line = line[2:].strip()
            line = line.replace('•', '').replace('·', '').strip()
            if line.startswith('#'):
                level = 0
                while level < len(line) and line[level] == '#': level += 1
                text = line[level:].strip().replace('*', '')
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE 
                r = p.add_run(text)
                r.bold = True
                r.font.size = Pt(12) 
                r.font.name = '黑体'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            else:
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE 
                add_runs_to_paragraph(p, line)

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue(), f"{title}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    except:
        return f"【{title}】\n\n{content}".encode('utf-8-sig'), f"{title}.txt", "text/plain"

@st.cache_data(ttl=2592000, show_spinner=False)
def get_ai_advice_for_student(grade, student_name, subject, weak_points, strong_points):
    if not client: return "⚠️ AI 尚未配置。"
    prompt = f"你是经验丰富的{grade}{subject}教师。学生 {student_name} 优势：{strong_points}。薄弱：{weak_points}。请结合具体知识点，写约300字的个性化鼓励和提分计划，给出具体的学习方法指导。"
    try:
        res = client.chat.completions.create(model="deepseek-chat", messages=[{"role": "system", "content": "你是专业AI导师。"}, {"role": "user", "content": prompt}])
        return res.choices[0].message.content
    except: return "AI 生成失败"

# ==============================================================================
# 🌐 左侧 SaaS 导航边栏 (动态响应登录状态)
# ==============================================================================
menu_sel = "成绩总览" # 默认值

with st.sidebar:
    st.markdown(f"<h2 style='margin-top:-20px; padding-bottom: 10px;'>🏫 英华学情查询</h2>", unsafe_allow_html=True)
    
    if st.session_state.logged_in_student:
        # 已登录：展示学生名片
        st.markdown(f"""
        <div style='background: linear-gradient(135deg, #1E293B 0%, #0F172A 100%); padding: 18px; border-radius: 12px; margin-bottom: 20px; border: 1px solid rgba(255,255,255,0.05); box-shadow: 0 4px 12px rgba(0,0,0,0.2);'>
            <div style='color: #F8FAFC; font-size: 18px; font-weight: 800; letter-spacing: 1px;'>👨‍🎓 {st.session_state.logged_in_student}</div>
            <div style='color: #94A3B8; font-size: 13px; margin-top: 6px;'>📚 {st.session_state.current_grade} · {st.session_state.logged_in_direction}</div>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("<p style='font-size: 13px; font-weight: bold; margin-bottom: -10px; color: #64748B;'>年级大区</p>", unsafe_allow_html=True)
        selected_grade = st.selectbox("隐藏标签1", ["高三", "高二", "高一"], index=["高三", "高二", "高一"].index(st.session_state.current_grade), label_visibility="collapsed")
        
        st.markdown("<div style='margin-top: 20px;'></div>", unsafe_allow_html=True)
        
        # 🔴 极简折叠导航
        with st.expander("📊 学情导航", expanded=True):
            menu_sel = option_menu(
                menu_title=None,
                options=["成绩总览", "历次追踪", "深度诊断"],
                icons=["grid-fill", "graph-up-arrow", "bullseye"],
                menu_icon="cast",
                default_index=0,
                styles={
                    "container": {"padding": "5px!important", "background-color": "#162032", "border-radius": "10px"},
                    "icon": {"color": "#64748B", "font-size": "15px"},
                    "nav-link": {"font-size": "14px", "text-align": "left", "margin":"0px 0", "color": "#64748B", "border-radius": "6px", "padding": "6px 15px"},
                    "nav-link-selected": {"background-color": "transparent", "color": "#F8FAFC", "font-weight": "bold"},
                    "nav-link:hover": {"background-color": "rgba(255,255,255,0.03)"}
                }
            )
            
        st.divider()
        st.button("🚪 退出", on_click=logout, use_container_width=True, type="secondary")
        
    else:
        # 未登录：仅展示年级选择
        st.markdown("<p style='font-size: 13px; font-weight: bold; margin-bottom: -10px; color: #64748B;'>选择查询年级</p>", unsafe_allow_html=True)
        selected_grade = st.selectbox("隐藏标签1", ["高三", "高二", "高一"], index=["高三", "高二", "高一"].index(st.session_state.current_grade), label_visibility="collapsed")

if selected_grade != st.session_state.current_grade:
    st.session_state.current_grade = selected_grade
    st.rerun()

# ==============================================================================
# 🚪 登录与主内容区
# ==============================================================================
if not st.session_state.logged_in_student:
    st.markdown(f"<h1 class='main-title'>🏫 英华学校【{selected_grade}】学情查询端</h1>", unsafe_allow_html=True)
    master_df, LATEST_EXAM, EXAMS_LIST = build_master_df(st.session_state.current_grade)
    
    if master_df is not None and not master_df.empty:
        top_p = master_df[master_df['方向'] == '物理方向'].sort_values('总分', ascending=False).head(5)['姓名'].tolist()
        top_h = master_df[master_df['方向'] == '历史方向'].sort_values('总分', ascending=False).head(5)['姓名'].tolist()
        str_p = f"🚀 理科前五名：{'、'.join(top_p)}" if top_p else ""
        str_h = f"🌟 文科前五名：{'、'.join(top_h)}" if top_h else ""
        banner_html = f"🎉 <b>【{LATEST_EXAM['name']}】成绩表彰光荣榜</b> 🏆<br>"
        if str_p: banner_html += f"<span style='font-size: 16px; color: #D97706;'>{str_p}</span>"
        if str_p and str_h: banner_html += "<br>"
        if str_h: banner_html += f"<span style='font-size: 16px; color: #D97706;'>{str_h}</span>"
        st.markdown(f'<div class="congrats-banner">{banner_html}</div>', unsafe_allow_html=True)
    
    col_left, col_mid, col_right = st.columns([1, 1.8, 1])
    with col_mid:
        with st.form("student_login"):
            st.markdown(f"<h2 style='text-align: center; color: #0F172A; margin-bottom: 30px; font-weight: 800;'>👨‍🎓 家长与学生统一入口</h2>", unsafe_allow_html=True)
            st.info("💡 提示：系统会自动根据您的学科分数识别文理方向。")
            name = st.text_input("👤 学生姓名", placeholder="请输入真实姓名")
            stu_id = st.text_input("🔢 考号/学号", placeholder="请输入准确考号")
            if st.form_submit_button("🔍 安全验证并查分", use_container_width=True):
                if name and stu_id:
                    if master_df is not None:
                        clean_n = clean_name(name)
                        clean_i = clean_str(stu_id)
                        match = master_df[(master_df['姓名'] == clean_n) & (master_df['考号'] == clean_i)]
                        if not match.empty:
                            st.session_state.logged_in_student = clean_n
                            st.session_state.logged_in_id = clean_i
                            st.session_state.logged_in_direction = match.iloc[0]['方向']
                            st.rerun()
                        else: st.error("❌ 未查询到成绩，请确认姓名和考号是否正确。")
                    else: st.warning("系统暂未配置该年级的考试数据。")
                else: st.error("⚠️ 请完整填写信息")
else:
    master_df, LATEST_EXAM, EXAMS_LIST = build_master_df(st.session_state.current_grade)
    stu_data = master_df[(master_df['姓名'] == st.session_state.logged_in_student) & (master_df['考号'] == st.session_state.logged_in_id)].iloc[0]
    
    # 顶部 Header 卡片
    st.markdown(f"""
    <div class="header-card">
        <h3 style="margin: 0; color: #0F172A; font-weight: 800;">❖ {menu_sel} <span style="font-size:16px; color:#94A3B8; font-weight:normal; margin-left: 10px;">/ {LATEST_EXAM['name']}</span></h3>
    </div>
    """, unsafe_allow_html=True)
    
    # ------------------ 模块一：成绩总览 ------------------
    if menu_sel == "成绩总览":
        k1, k2, k3, k4, k5 = st.columns(5)
        k1.metric("学生姓名", stu_data['姓名'])
        k2.metric("文理方向", stu_data['方向'])
        k3.metric("总分实考", f"{stu_data['总分']} 分")
        k4.metric("班级名次", f"第 {stu_data['总分班级排名']} 名")
        k5.metric("年级名次", f"第 {stu_data['总分年级排名']} 名")
        
        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown("<p style='font-size: 18px; font-weight: bold; color: #0F172A;'>📊 各科得分雷达扫描</p>", unsafe_allow_html=True)
        
        subs = ['语文','数学','英语','物理','化学','生物','历史','政治','地理']
        valid_subs = [s for s in subs if s in stu_data and stu_data[s] > 0]
        
        if valid_subs:
            chart_data = pd.DataFrame({"科目": valid_subs, "得分": [stu_data[s] for s in valid_subs]})
            col_bar, col_radar = st.columns(2)
            with col_bar:
                with st.container(border=True):
                    fig1 = px.bar(chart_data, x='科目', y='得分', text_auto=True, color='科目', color_discrete_sequence=px.colors.qualitative.Pastel)
                    fig1.update_traces(textposition='outside', width=0.35, textfont_size=13, marker_line_width=0)
                    y_max = chart_data['得分'].max() * 1.15
                    fig1.update_layout(showlegend=False, margin=dict(t=40, b=20, l=20, r=20), paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(248, 250, 252, 0.5)', yaxis_range=[0, y_max], yaxis=dict(showgrid=True, gridcolor='#F1F5F9'), dragmode=False)
                    st.plotly_chart(fig1, use_container_width=True, config=CHART_CONFIG)
            with col_radar:
                with st.container(border=True):
                    fig2 = px.line_polar(chart_data, r='得分', theta='科目', line_close=True)
                    fig2.update_traces(fill='toself', line_color='#3B82F6')
                    fig2.update_layout(margin=dict(t=40, b=20, l=40, r=40), paper_bgcolor='rgba(0,0,0,0)', polar=dict(radialaxis=dict(visible=True, showline=False)), dragmode=False)
                    st.plotly_chart(fig2, use_container_width=True, config=CHART_CONFIG)
            
    # ------------------ 模块二：历次追踪 ------------------
    elif menu_sel == "历次追踪":
        history_records = []
        with st.spinner("正在汇聚历次成绩轨迹..."):
            for exam in EXAMS_LIST:
                exam_df = None
                dfs_sub = []
                subs_hist = ['语文','数学','英语','物理','化学','生物','历史','政治','地理']
                for sh in subs_hist:
                    u = exam.get(sh)
                    if u:
                        d_sub = load_data(u, header_lines=[0,1,2])
                        if d_sub is not None:
                            n_c, i_c = None, None
                            for col in d_sub.columns:
                                cstr = str(col[0]) if isinstance(col, tuple) else str(col)
                                if '姓名' in cstr: n_c = col
                                elif '考号' in cstr or '学号' in cstr: i_c = col
                            if n_c and i_c:
                                r_ls = []
                                for _, row in d_sub.iterrows():
                                    if clean_str(row[i_c]) == st.session_state.logged_in_id:
                                        tot = 0
                                        for c in d_sub.columns:
                                            if c in [n_c, i_c]: continue
                                            if '总分' in str(c) or '排名' in str(c): continue
                                            try: 
                                                v = float(row[c])
                                                if pd.notna(v): tot += v
                                            except: pass
                                        r_ls.append({'考号': st.session_state.logged_in_id, sh: round(tot, 1)})
                                if r_ls: dfs_sub.append(pd.DataFrame(r_ls))
                if dfs_sub:
                    exam_master = functools.reduce(lambda l, r: pd.merge(l, r, on=['考号'], how='outer'), dfs_sub)
                    p_subs = [s for s in subs_hist if s in exam_master.columns]
                    exam_master['总分'] = exam_master[p_subs].fillna(0).sum(axis=1).round(1)
                    row = exam_master.iloc[0]
                    rec = {"考试名称": exam['name']}
                    if '总分' in row: rec['总分'] = float(row['总分'])
                    for s in subs_hist:
                        if s in row and pd.notna(row[s]): rec[s] = float(row[s])
                    history_records.append(rec)
        
        if history_records:
            df_trend = pd.DataFrame(history_records)
            with st.container(border=True):
                col_t1, col_t2 = st.columns(2)
                with col_t1:
                    fig_score = px.line(df_trend, x="考试名称", y="总分", markers=True, title="📈 总分走势图", line_shape="spline")
                    fig_score.update_traces(line_color="#3B82F6", marker=dict(size=10))
                    fig_score.update_layout(dragmode=False, paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)', yaxis=dict(showgrid=True, gridcolor='#F1F5F9'))
                    st.plotly_chart(fig_score, use_container_width=True, config=CHART_CONFIG)
                with col_t2:
                    if "年级排名" in df_trend.columns: # 临时处理历史没有排名的问题
                        pass
                    else:
                        st.info("单科波动透视")
                        avail_hist_subs = [s for s in ['语文','数学','英语','物理','化学','生物','历史','政治','地理'] if s in df_trend.columns]
                        if avail_hist_subs:
                            sel_hist_sub = st.selectbox("选择科目：", avail_hist_subs, key="hist_sub", label_visibility="collapsed")
                            fig3 = px.line(df_trend, x="考试名称", y=sel_hist_sub, markers=True, title=f"📉 {sel_hist_sub} 单科走势", line_shape="spline")
                            fig3.update_traces(line_color="#10B981", marker=dict(size=10))
                            fig3.update_layout(dragmode=False, paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)', yaxis=dict(showgrid=True, gridcolor='#F1F5F9'))
                            st.plotly_chart(fig3, use_container_width=True, config=CHART_CONFIG)
        else: st.info("暂未抓取到您的历史轨迹。")

    # ------------------ 模块三：深度诊断 ------------------
    elif menu_sel == "深度诊断":
        avail_subs = [s for s in ['语文','数学','英语','物理','化学','生物','历史','政治','地理'] if s in stu_data and stu_data[s] > 0 and LATEST_EXAM.get(s)]
        if not avail_subs: st.info("暂未配置您所考科目的详细题库数据。")
        else:
            st.markdown("<p style='font-size: 14px; font-weight: bold; color: #64748B; margin-bottom: -10px;'>⚙️ 选择诊断科目</p>", unsafe_allow_html=True)
            c_sel, _ = st.columns([1, 4])
            sel_sub = c_sel.selectbox("隐藏标签", avail_subs, label_visibility="collapsed")
            st.markdown("<br>", unsafe_allow_html=True)
            
            df_diag = load_data(LATEST_EXAM[sel_sub], header_lines=[0, 1, 2])
            if df_diag is not None:
                name_c, id_c, cls_c = None, None, None
                for col in df_diag.columns:
                    cstr = str(col[0]) if isinstance(col, tuple) else str(col)
                    if '姓名' in cstr: name_c = col
                    elif '考号' in cstr or '学号' in cstr: id_c = col
                    elif '班级' in cstr: cls_c = col
                if name_c and id_c:
                    found_idx = -1
                    for idx, row in df_diag.iterrows():
                        if clean_name(row[name_c]) == st.session_state.logged_in_student and clean_str(row[id_c]) == st.session_state.logged_in_id: 
                            found_idx = idx; break
                    if found_idx != -1:
                        knowledge_map = {} 
                        for col in df_diag.columns:
                            if col in [name_c, id_c, cls_c]: continue
                            cstr = str(col[0]) if isinstance(col, tuple) else str(col)
                            if '总分' in cstr or '排名' in cstr: continue
                            q_name = str(col[0]).strip() if isinstance(col, tuple) else str(col).strip()
                            k_point = str(col[1]).strip() if isinstance(col, tuple) and len(col) > 1 else q_name
                            if k_point == "" or k_point.startswith("Unnamed"): k_point = q_name
                            try: full = float(col[2]) if isinstance(col, tuple) and len(col) > 2 else 0
                            except: full = 0
                            if full <= 0:
                                try: full = float(pd.to_numeric(df_diag[col], errors='coerce').max())
                                except: full = 0
                            if full <= 0: continue
                            if k_point not in knowledge_map: knowledge_map[k_point] = {'my': 0, 'full': 0, 'class_total': 0}
                            try: my_s = float(df_diag.iloc[found_idx][col])
                            except: my_s = 0
                            class_s = pd.to_numeric(df_diag[col], errors='coerce').mean()
                            knowledge_map[k_point]['my'] += my_s
                            knowledge_map[k_point]['full'] += full
                            knowledge_map[k_point]['class_total'] += class_s
                        
                        k_data, weak_points_list, strong_points_list = [], [], []
                        for kp, val in knowledge_map.items():
                            my_rate = round((val['my']/val['full'])*100, 1) if val['full']>0 else 0
                            avg_rate = round((val['class_total']/val['full'])*100, 1) if val['full']>0 else 0
                            k_data.append({'知识点': kp, '我的掌握率': my_rate, '班级平均': avg_rate})
                            if my_rate < avg_rate: weak_points_list.append(kp)
                            else: strong_points_list.append(kp)
                        
                        df_kp = pd.DataFrame(k_data)
                        if not df_kp.empty:
                            with st.container(border=True):
                                c_chart, c_text = st.columns([1.2, 1])
                                with c_chart:
                                    fig = go.Figure()
                                    cats = df_kp['知识点'].tolist() + [df_kp['知识点'].tolist()[0]]
                                    mys = df_kp['我的掌握率'].tolist() + [df_kp['我的掌握率'].tolist()[0]]
                                    avgs = df_kp['班级平均'].tolist() + [df_kp['班级平均'].tolist()[0]]
                                    fig.add_trace(go.Scatterpolar(r=avgs, theta=cats, fill='toself', name='班级平均', line_color='#CBD5E1'))
                                    fig.add_trace(go.Scatterpolar(r=mys, theta=cats, fill='toself', name='我的掌握', line_color='#3B82F6'))
                                    fig.update_layout(polar=dict(radialaxis=dict(visible=True, range=[0, 100])), paper_bgcolor='rgba(0,0,0,0)', dragmode=False)
                                    st.plotly_chart(fig, use_container_width=True, config=CHART_CONFIG)
                                with c_text:
                                    st.markdown("<p style='font-size: 16px; font-weight: bold; color: #0F172A;'>🩺 薄弱知识点定位</p>", unsafe_allow_html=True)
                                    if weak_points_list:
                                        for row in k_data:
                                            if row['知识点'] in weak_points_list:
                                                st.markdown(f"<p style='color: #EF4444; font-size: 14px;'>▪ <b>{row['知识点']}</b> (落后班级 {row['班级平均'] - row['我的掌握率']:.1f}%)</p>", unsafe_allow_html=True)
                                    else: st.success("🎉 所有知识点均达标！")
                            
                            st.divider()
                            
                            ai_state_key = f"ai_stu_{st.session_state.current_grade}_{st.session_state.logged_in_id}_{sel_sub}"
                            saved_list_key = f"saved_ai_stu_list_{st.session_state.current_grade}_{st.session_state.logged_in_id}_{sel_sub}"
                            if saved_list_key not in st.session_state: st.session_state[saved_list_key] = []

                            if AI_API_KEY:
                                if st.button(f"✨ 召唤 AI 导师生成专属提分计划", type="primary"):
                                    with st.spinner("AI 导师正在分析数据..."):
                                        w_str = "、".join(weak_points_list) if weak_points_list else "无"
                                        s_str = "、".join(strong_points_list) if strong_points_list else "无"
                                        ai_reply = get_ai_advice_for_student(st.session_state.current_grade, st.session_state.logged_in_student, sel_sub, w_str, s_str)
                                        st.session_state[ai_state_key] = ai_reply

                                if ai_state_key in st.session_state:
                                    saved_reply = st.session_state[ai_state_key]
                                    st.markdown(f"<div class='ai-box'><b>👨‍🏫 导师寄语：</b><br><br>{saved_reply}</div><br>", unsafe_allow_html=True)
                                    
                                    doc_title = f"【{LATEST_EXAM['name']}】{st.session_state.logged_in_student}_{sel_sub}_提分计划"
                                    t_c1, t_c2, t_c3 = st.columns([1.5, 1, 1])
                                    with t_c1:
                                        if st.button("📌 存入本机档案库"):
                                            st.session_state[saved_list_key].insert(0, saved_reply)
                                            st.toast("✅ 已成功存入！")
                                    with t_c2:
                                        export_fmt = st.selectbox("隐藏", ["Word文档 (自动精排版)", "TXT纯文本"], label_visibility="collapsed", key="fmt_stu")
                                    with t_c3:
                                        if "Word" in export_fmt:
                                            file_data, file_name, mime_type = generate_ai_doc(doc_title, saved_reply)
                                        else:
                                            file_data = saved_reply.encode('utf-8-sig')
                                            file_name = f"{doc_title}.txt"
                                            mime_type = "text/plain"
                                        st.download_button(label="📥 导出至电脑", data=file_data, file_name=file_name, mime=mime_type, type="primary")

                                if st.session_state[saved_list_key]:
                                    with st.expander(f"📂 历史暂存报告 (共 {len(st.session_state[saved_list_key])} 份)"):
                                        for idx, old_rep in enumerate(st.session_state[saved_list_key]):
                                            st.markdown(f"**🔖 版本 {len(st.session_state[saved_list_key]) - idx}**")
                                            st.markdown(old_rep)
                                            st.divider()