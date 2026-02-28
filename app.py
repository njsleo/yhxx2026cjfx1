import streamlit as st
import pandas as pd
import plotly.graph_objects as go
import plotly.express as px
from streamlit_option_menu import option_menu
import openai
import os
import functools
import io
import re

# ==============================================================================
# 1. 页面基础配置 
# ==============================================================================
st.set_page_config(page_title="英华学校高中部考试学情智能分析", layout="wide", page_icon="🏫", initial_sidebar_state="collapsed")

# ==============================================================================
# 🔐 安全配置与【时间胶囊】动态数据源读取
# ==============================================================================
try:
    ADMIN_PASSWORD = st.secrets["ADMIN_PWD"]
    HOMEROOM_PASSWORD = st.secrets.get("HOMEROOM_PWD", ADMIN_PASSWORD) 
    TEACHER_PASSWORD = st.secrets.get("TEACHER_PWD", ADMIN_PASSWORD)
    URL_TEACHER_ROSTER = st.secrets.get("URL_TEACHER_ROSTER", "") 
    AI_API_KEY = st.secrets.get("DEEPSEEK_API_KEY", "")
    
    EXAMS = []
    for i in range(1, 21):
        name = st.secrets.get(f"EXAM_NAME_{i}")
        if name:
            EXAMS.append({
                "name": name,
                "语文": st.secrets.get(f"URL_CHINESE_{i}", ""),
                "数学": st.secrets.get(f"URL_MATH_{i}", ""),
                "英语": st.secrets.get(f"URL_ENGLISH_{i}", ""),
                "物理": st.secrets.get(f"URL_PHYSICS_{i}", ""),
                "化学": st.secrets.get(f"URL_CHEMISTRY_{i}", ""),
                "生物": st.secrets.get(f"URL_BIOLOGY_{i}", ""),
                "历史": st.secrets.get(f"URL_HISTORY_{i}", ""),
                "政治": st.secrets.get(f"URL_POLITICS_{i}", ""),
                "地理": st.secrets.get(f"URL_GEOGRAPHY_{i}", "")
            })
            
    LATEST_EXAM_IDX = len(EXAMS) - 1 if EXAMS else -1
    LATEST_EXAM = EXAMS[-1] if EXAMS else None

except Exception as e:
    st.error("⚠️ 系统配置读取失败，请检查 Streamlit 后台的 Secrets。")
    st.stop()

if AI_API_KEY: client = openai.OpenAI(api_key=AI_API_KEY, base_url="https://api.deepseek.com")
else: client = None

# ==============================================================================
# 🛠️ 核心引擎：数据净化、班级名统一、总分聚合及【双排名计算】
# ==============================================================================
def clean_str(val):
    if pd.isna(val): return ""
    v = str(val).strip()
    if v.endswith('.0'): v = v[:-2]
    return v

def clean_name(val):
    if pd.isna(val): return ""
    return str(val).replace(" ", "").strip()

def normalize_class_name(c):
    if pd.isna(c): return ""
    c = str(c).replace(" ", "").strip()
    mapping = {'1':'一', '2':'二', '3':'三', '4':'四', '5':'五', '6':'六', '7':'七', '8':'八', '9':'九', '0':'零'}
    for k, v in mapping.items():
        c = c.replace(k, v)
    c = c.replace("高三", "").replace("高二", "").replace("高一", "").replace("年级", "").replace("()", "").replace("（）", "")
    if not c.endswith("班"): c += "班"
    return c

@st.cache_data(ttl=600)
def load_data(url, header_lines=0):
    if not url or not url.strip(): return None
    try: return pd.read_csv(url, header=header_lines, on_bad_lines='skip')
    except: return None

@st.cache_data(ttl=600, show_spinner=False)
def build_master_df(exam_idx):
    if exam_idx < 0 or exam_idx >= len(EXAMS): return None
    exam = EXAMS[exam_idx]
    dfs = []
    subs = ['语文','数学','英语','物理','化学','生物','历史','政治','地理']
    
    for sub in subs:
        url = exam.get(sub)
        if url:
            df_sub = load_data(url, header_lines=[0,1,2])
            if df_sub is not None:
                name_c, id_c, cls_c = None, None, None
                for col in df_sub.columns:
                    cstr = str(col[0]) if isinstance(col, tuple) else str(col)
                    if '姓名' in cstr: name_c = col
                    elif '考号' in cstr or '学号' in cstr: id_c = col
                    elif '班级' in cstr: cls_c = col
                    
                if name_c and id_c:
                    res = []
                    for _, row in df_sub.iterrows():
                        tot = 0
                        for c in df_sub.columns:
                            if c in [name_c, id_c, cls_c]: continue
                            cstr = str(c[0]) if isinstance(c, tuple) else str(c)
                            if '总分' in cstr or '排名' in cstr: continue
                            try: 
                                val = float(row[c])
                                if pd.notna(val): tot += val
                            except: pass
                                
                        s_name = clean_name(row[name_c])
                        s_id = clean_str(row[id_c])
                        s_cls = normalize_class_name(row[cls_c]) if cls_c else "未分班"
                        
                        if s_id: res.append({'姓名': s_name, '考号': s_id, '班级': s_cls, sub: round(tot, 1)})
                    if res: dfs.append(pd.DataFrame(res))
    
    if not dfs: return None
    master = functools.reduce(lambda l, r: pd.merge(l, r, on=['姓名','考号','班级'], how='outer'), dfs)
    present_subs = [s for s in subs if s in master.columns]
    master[present_subs] = master[present_subs].fillna(0)
    master['总分'] = master[present_subs].sum(axis=1).round(1)
    
    def get_dir(r):
        if r.get('物理', 0) > 0: return "物理方向"
        if r.get('历史', 0) > 0: return "历史方向"
        return "综合方向"
    master['方向'] = master.apply(get_dir, axis=1)
    
    master['班级排名'] = master.groupby(['班级', '方向'])['总分'].rank(ascending=False, method='min').fillna(0).astype(int)
    master['年级排名'] = master.groupby(['方向'])['总分'].rank(ascending=False, method='min').fillna(0).astype(int)
    return master

# ==============================================================================
# 🎨 网页版原生 HTML 高颜值斑马线表格渲染器
# ==============================================================================
def render_html_table(df):
    html = """
    <div style="width: 100%; overflow-x: auto; margin-bottom: 25px; border-radius: 10px; box-shadow: 0 4px 12px rgba(0,0,0,0.08);">
    <table style="width: 100%; border-collapse: collapse; font-size: 16px; text-align: center; font-family: 'Helvetica Neue', Arial, sans-serif;">
    """
    html += "<tr>"
    for col in df.columns:
        html += f"<th style='background-color: #0068C9; color: white; padding: 18px 12px; border: 1px solid #e1e4e8; white-space: nowrap; font-weight: bold;'>{col}</th>"
    html += "</tr>"
    
    for i, row in df.iterrows():
        bg_color = "#F8FAFC" if i % 2 == 0 else "#FFFFFF"
        html += f"<tr style='background-color: {bg_color}; transition: background-color 0.2s;' onmouseover=\"this.style.backgroundColor='#E6F3FF'\" onmouseout=\"this.style.backgroundColor='{bg_color}'\">"
        for col in df.columns:
            val = row[col]
            if isinstance(val, float): val = f"{val:.1f}"
            html += f"<td style='padding: 16px 12px; border: 1px solid #e1e4e8; color: #2c3e50;'>{val}</td>"
        html += "</tr>"
    html += "</table></div>"
    st.markdown(html, unsafe_allow_html=True)

# ==============================================================================
# 📥 导出神器 1：Excel 成绩表生成器
# ==============================================================================
def generate_excel_download(df, filename_prefix, title_text):
    try:
        from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
        from openpyxl.utils import get_column_letter
        
        buffer = io.BytesIO()
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='成绩明细', startrow=1)
            worksheet = writer.sheets['成绩明细']
            num_cols = len(df.columns)
            
            worksheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=num_cols)
            title_cell = worksheet.cell(row=1, column=1, value=title_text)
            title_cell.font = Font(size=20, bold=True, color="FFFFFFFF") 
            title_cell.fill = PatternFill(start_color="FF0068C9", end_color="FF0068C9", fill_type="solid") 
            title_cell.alignment = Alignment(horizontal="center", vertical="center")
            worksheet.row_dimensions[1].height = 45 
            
            header_fill = PatternFill(start_color="FF4A90E2", end_color="FF4A90E2", fill_type="solid")
            header_font = Font(bold=True, color="FFFFFFFF", size=12)
            even_fill = PatternFill(start_color="FFF8FAFC", end_color="FFF8FAFC", fill_type="solid")
            odd_fill = PatternFill(start_color="FFFFFFFF", end_color="FFFFFFFF", fill_type="solid")
            thin_border = Border(left=Side(style='thin', color='FFDDDDDD'), right=Side(style='thin', color='FFDDDDDD'), 
                                 top=Side(style='thin', color='FFDDDDDD'), bottom=Side(style='thin', color='FFDDDDDD'))
            
            for col_idx in range(1, num_cols + 1):
                col_letter = get_column_letter(col_idx)
                max_len = sum(2 if ord(c)>127 else 1 for c in str(df.columns[col_idx-1]))
                
                for row_idx in range(2, len(df) + 3):
                    cell = worksheet.cell(row=row_idx, column=col_idx)
                    cell.alignment = Alignment(horizontal="center", vertical="center")
                    cell.border = thin_border
                    
                    if row_idx == 2:
                        cell.fill = header_fill
                        cell.font = header_font
                    else:
                        cell.fill = even_fill if row_idx % 2 == 0 else odd_fill
                        cell.font = Font(size=11)
                        
                    val_str = str(cell.value) if cell.value is not None else ""
                    val_len = sum(2 if ord(c)>127 else 1 for c in val_str)
                    if val_len > max_len: max_len = val_len
                        
                worksheet.column_dimensions[col_letter].width = max_len + 4

            for row_idx in range(2, len(df) + 3):
                worksheet.row_dimensions[row_idx].height = 25

        return buffer.getvalue(), f"{filename_prefix}.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    except Exception as e:
        return df.to_csv(index=False).encode('utf-8-sig'), f"{filename_prefix}.csv", "text/csv"

# ==============================================================================
# 📥 导出神器 2：AI 报告生成器 (优先Word，保底TXT)
# ==============================================================================
def generate_ai_doc(title, content):
    try:
        import docx
        from docx.shared import Pt
        doc = docx.Document()
        doc.add_heading(title, level=1)
        for para in content.split('\n'):
            if para.strip():
                p = doc.add_paragraph(para.strip())
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue(), f"{title}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    except:
        # 保底降级为 TXT
        text_content = f"【{title}】\n\n{content}"
        return text_content.encode('utf-8-sig'), f"{title}.txt", "text/plain"

# ==============================================================================
# 🧠 AI 导师功能定义
# ==============================================================================
@st.cache_data(ttl=2592000, show_spinner=False)
def get_ai_advice_for_student(student_name, subject, weak_points, strong_points):
    if not client: return "⚠️ AI 尚未配置。"
    prompt = f"你是经验丰富的高中{subject}教师。学生 {student_name} 优势：{strong_points}。薄弱：{weak_points}。写约300字的个性化提分计划。"
    try:
        res = client.chat.completions.create(model="deepseek-chat", messages=[{"role": "system", "content": "你是专业AI导师。"}, {"role": "user", "content": prompt}])
        return res.choices[0].message.content
    except: return "AI 生成失败"

@st.cache_data(ttl=2592000, show_spinner=False)
def get_ai_advice_for_teacher(subject, weak_points_list):
    if not client: return "⚠️ AI 尚未配置。"
    prompt = f"你是教研员。高三年级{subject}薄弱点是：{weak_points_list}。给老师们写约300字的讲评教研建议。"
    try:
        res = client.chat.completions.create(model="deepseek-chat", messages=[{"role": "system", "content": "你是教研专家AI。"}, {"role": "user", "content": prompt}])
        return res.choices[0].message.content
    except: return "AI 生成失败"

# ==============================================================================
# --- 状态与 CSS 样式 ---
# ==============================================================================
if 'logged_in_student' not in st.session_state: st.session_state.logged_in_student = None
if 'logged_in_id' not in st.session_state: st.session_state.logged_in_id = None
if 'logged_in_direction' not in st.session_state: st.session_state.logged_in_direction = None
if 'teacher_role' not in st.session_state: st.session_state.teacher_role = None 
if 'teacher_name' not in st.session_state: st.session_state.teacher_name = None
if 'teacher_subject' not in st.session_state: st.session_state.teacher_subject = None
if 'teacher_classes' not in st.session_state: st.session_state.teacher_classes = []

def logout():
    for key in st.session_state.keys(): st.session_state[key] = None
    st.rerun()

st.markdown("""
<style>
    #MainMenu {visibility: hidden;} header {visibility: hidden;} footer {visibility: hidden;}
    .block-container { padding-top: 1rem !important; padding-bottom: 2rem !important; }
    .stApp { background-color: #f4f7f9; }
    div[data-testid="stMetric"] { background-color: #ffffff; border-radius: 12px; padding: 20px; box-shadow: 0 4px 10px rgba(0,0,0,0.03); border: 1px solid #ebeef5; text-align: center; }
    div[data-testid="stForm"] { background-color: #ffffff; padding: 30px; border-radius: 15px; box-shadow: 0 8px 20px rgba(0,0,0,0.05); border: none; }
    div[data-testid="stFormSubmitButton"] > button { background-color: #0068C9; color: white; font-weight: bold; border-radius: 8px; border: none; padding: 10px 0; }
    .congrats-banner { background: linear-gradient(90deg, #FFFBEB, #FFF7ED); border: 2px solid #FCD34D; color: #92400E; padding: 12px 20px; border-radius: 12px; text-align: center; font-size: 18px; font-weight: bold; margin-bottom: 25px; box-shadow: 0 4px 12px rgba(252, 211, 77, 0.2); line-height: 1.6; }
    .main-title { text-align: center; color: #1E3A8A; font-size: 28px; font-weight: 800; margin-bottom: 15px; }
    .ai-box { background: linear-gradient(135deg, #f0f7ff 0%, #e6f3ff 100%); border-left: 5px solid #0068C9; padding: 20px; border-radius: 8px; font-size: 15px; color: #333; line-height: 1.8;}
</style>
""", unsafe_allow_html=True)

CHART_CONFIG = {'displayModeBar': False, 'scrollZoom': False}

selected_nav = option_menu(
    menu_title=None, options=["成绩总览", "历次追踪", "深度诊断", "教师后台"], 
    icons=["clipboard-data", "graph-up", "bullseye", "person-badge"], menu_icon="cast", default_index=0, orientation="horizontal",
    styles={ "container": {"padding": "5px", "background-color": "#ffffff", "border-radius": "12px", "box-shadow": "0 4px 15px rgba(0,0,0,0.08)", "margin-bottom": "30px", "position": "sticky", "top": "15px", "z-index": "9999"}, "nav-link-selected": {"background-color": "#0068C9", "color": "white", "font-weight": "bold"} }
)

# ==============================================================================
# 🚀 页面逻辑：学生端
# ==============================================================================
if selected_nav in ["成绩总览", "历次追踪", "深度诊断"]:
    
    if not st.session_state.logged_in_student:
        st.markdown("<h1 class='main-title'>🏫 英华学校高中部考试学情智能分析系统</h1>", unsafe_allow_html=True)
        
        latest_master = build_master_df(LATEST_EXAM_IDX)
        if latest_master is not None and not latest_master.empty:
            top_p = latest_master[latest_master['方向'] == '物理方向'].sort_values('总分', ascending=False).head(5)['姓名'].tolist()
            top_h = latest_master[latest_master['方向'] == '历史方向'].sort_values('总分', ascending=False).head(5)['姓名'].tolist()
            str_p = f"🚀 理科前五：{'、'.join(top_p)}" if top_p else ""
            str_h = f"🌟 文科前五：{'、'.join(top_h)}" if top_h else ""
            
            banner_html = f"🎉 <b>【{LATEST_EXAM['name']}】成绩表彰光荣榜</b> 🏆<br>"
            if str_p: banner_html += f"<span style='font-size: 16px; color: #D97706;'>{str_p}</span>"
            if str_p and str_h: banner_html += "<br>"
            if str_h: banner_html += f"<span style='font-size: 16px; color: #D97706;'>{str_h}</span>"
            st.markdown(f'<div class="congrats-banner">{banner_html}</div>', unsafe_allow_html=True)
        
        col_left, col_mid, col_right = st.columns([1, 1.8, 1])
        with col_left:
            st.markdown("<br><br>", unsafe_allow_html=True)
            if os.path.exists("panda.gif"): st.image("panda.gif", use_container_width=True)
        with col_mid:
            with st.form("student_login"):
                st.markdown("<h3 style='text-align: center; color: #555;'>👨‍🎓 学生/家长登录入口</h3><br>", unsafe_allow_html=True)
                st.info("💡 提示：系统会自动根据您的学科分数识别文理方向。")
                name = st.text_input("👤 学生姓名", placeholder="请输入真实姓名")
                stu_id = st.text_input("🔢 考号/学号", placeholder="请输入准确考号")
                if st.form_submit_button("🔍 立即查分", use_container_width=True):
                    if name and stu_id:
                        if latest_master is not None:
                            clean_n = clean_name(name)
                            clean_i = clean_str(stu_id)
                            match = latest_master[(latest_master['姓名'] == clean_n) & (latest_master['考号'] == clean_i)]
                            if not match.empty:
                                st.session_state.logged_in_student = clean_n
                                st.session_state.logged_in_id = clean_i
                                st.session_state.logged_in_direction = match.iloc[0]['方向']
                                st.rerun()
                            else: st.error("❌ 未查询到成绩，请确认姓名和考号是否正确。")
                        else: st.warning("系统暂未配置考试数据。")
                    else: st.error("⚠️ 请完整填写信息")
        with col_right:
            st.markdown("<br><br>", unsafe_allow_html=True)
            if os.path.exists("star.gif"): st.image("star.gif", use_container_width=True)
    
    else:
        c1, c2 = st.columns([4, 1])
        c1.markdown(f"**当前用户：** {st.session_state.logged_in_student} | **系统判定方向：** {st.session_state.logged_in_direction}")
        if c2.button("🚪 退出登录", use_container_width=True): logout()
        st.divider()

        master_df = build_master_df(LATEST_EXAM_IDX)
        stu_data = master_df[(master_df['姓名'] == st.session_state.logged_in_student) & (master_df['考号'] == st.session_state.logged_in_id)].iloc[0]
        
        # --- 模块 1: 成绩总览 ---
        if selected_nav == "成绩总览":
            st.markdown(f"### 🏆 【{LATEST_EXAM['name']}】成绩概览")
            k1, k2, k3, k4, k5 = st.columns(5)
            k1.metric("姓名", stu_data['姓名'])
            k2.metric("方向", stu_data['方向'])
            k3.metric("总分", f"{stu_data['总分']}")
            k4.metric("班级名次", f"第 {stu_data['总分班级排名']} 名")
            k5.metric("年级名次", f"第 {stu_data['总分年级排名']} 名")
            
            st.markdown("<br>", unsafe_allow_html=True)
            st.markdown("### 📊 各科得分对比")
            
            subs = ['语文','数学','英语','物理','化学','生物','历史','政治','地理']
            valid_subs = [s for s in subs if s in stu_data and stu_data[s] > 0]
            
            if valid_subs:
                chart_data = pd.DataFrame({"科目": valid_subs, "得分": [stu_data[s] for s in valid_subs]})
                col_bar, col_radar = st.columns(2)
                with col_bar:
                    fig1 = px.bar(chart_data, x='科目', y='得分', text_auto=True, color='科目')
                    fig1.update_layout(showlegend=False, margin=dict(t=40, b=20, l=20, r=20), paper_bgcolor='rgba(0,0,0,0)', dragmode=False)
                    st.plotly_chart(fig1, use_container_width=True, config=CHART_CONFIG)
                with col_radar:
                    fig2 = px.line_polar(chart_data, r='得分', theta='科目', line_close=True)
                    fig2.update_traces(fill='toself', line_color='#0068C9')
                    fig2.update_layout(margin=dict(t=40, b=20, l=40, r=40), paper_bgcolor='rgba(0,0,0,0)', dragmode=False)
                    st.plotly_chart(fig2, use_container_width=True, config=CHART_CONFIG)
            
        # --- 模块 2: 历次追踪 ---
        elif selected_nav == "历次追踪":
            st.markdown(f"### 📈 历次考试波动轨迹")
            history_records = []
            with st.spinner("正在云端汇聚历次成绩轨迹..."):
                for i, exam in enumerate(EXAMS):
                    m_df = build_master_df(i)
                    if m_df is not None and not m_df.empty:
                        stu_h = m_df[(m_df['姓名'] == st.session_state.logged_in_student) & (m_df['考号'] == st.session_state.logged_in_id)]
                        if not stu_h.empty:
                            history_records.append({ "考试名称": exam['name'], "总分": float(stu_h.iloc[0]['总分']), "年级排名": int(stu_h.iloc[0]['总分年级排名']) })
            
            if history_records:
                df_trend = pd.DataFrame(history_records)
                col_t1, col_t2 = st.columns(2)
                with col_t1:
                    fig_score = px.line(df_trend, x="考试名称", y="总分", markers=True, title="总分走势图", line_shape="spline")
                    fig_score.update_traces(line_color="#FF4B4B", marker=dict(size=10))
                    fig_score.update_layout(dragmode=False)
                    st.plotly_chart(fig_score, use_container_width=True, config=CHART_CONFIG)
                with col_t2:
                    fig_rank = px.line(df_trend, x="考试名称", y="年级排名", markers=True, title="总分年级排名走势 (向下代表进步)", line_shape="spline")
                    fig_rank.update_traces(line_color="#0068C9", marker=dict(size=10))
                    fig_rank.update_yaxes(autorange="reversed")
                    fig_rank.update_layout(dragmode=False)
                    st.plotly_chart(fig_rank, use_container_width=True, config=CHART_CONFIG)
            else: st.info("暂未抓取到您的历史轨迹。")

        # --- 模块 3: 深度诊断 ---
        elif selected_nav == "深度诊断":
            avail_subs = [s for s in ['语文','数学','英语','物理','化学','生物','历史','政治','地理'] if s in stu_data and stu_data[s] > 0 and LATEST_EXAM.get(s)]
            if not avail_subs: st.info("暂未配置您所考科目的详细题库数据。")
            else:
                sel_sub = st.selectbox("👇 选择诊断科目", avail_subs)
                df_diag = load_data(LATEST_EXAM[sel_sub], header_lines=[0, 1, 2])
                if df_diag is not None:
                    name_c, id_c, cls_c = None, None, None
                    for col in df_diag.columns:
                        cstr = str(col[0]) if isinstance(col, tuple) else str(col)
                        if '姓名' in cstr: name_c = col
                        elif '考号' in cstr or '学号' in cstr: id_c = col
                        elif '班级' in cstr: cls_c = col
                        
                    if name_c and id_c:
                        found_idx = -1
                        for idx, row in df_diag.iterrows():
                            if clean_name(row[name_c]) == st.session_state.logged_in_student and clean_str(row[id_c]) == st.session_state.logged_in_id: 
                                found_idx = idx; break
                                
                        if found_idx != -1:
                            knowledge_map = {} 
                            for col in df_diag.columns:
                                if col in [name_c, id_c, cls_c]: continue
                                cstr = str(col[0]) if isinstance(col, tuple) else str(col)
                                if '总分' in cstr or '排名' in cstr: continue
                                
                                q_name = str(col[0]).strip() if isinstance(col, tuple) else str(col).strip()
                                k_point = str(col[1]).strip() if isinstance(col, tuple) and len(col) > 1 else q_name
                                if k_point == "" or k_point.startswith("Unnamed"): k_point = q_name
                                
                                try: full = float(col[2]) if isinstance(col, tuple) and len(col) > 2 else 0
                                except: full = 0
                                
                                if full <= 0:
                                    try: full = float(pd.to_numeric(df_diag[col], errors='coerce').max())
                                    except: full = 0
                                
                                if full <= 0: continue
                                
                                if k_point not in knowledge_map: knowledge_map[k_point] = {'my': 0, 'full': 0, 'class_total': 0}
                                try: my_s = float(df_diag.iloc[found_idx][col])
                                except: my_s = 0
                                class_s = pd.to_numeric(df_diag[col], errors='coerce').mean()
                                
                                knowledge_map[k_point]['my'] += my_s
                                knowledge_map[k_point]['full'] += full
                                knowledge_map[k_point]['class_total'] += class_s
                            
                            k_data, weak_points_list, strong_points_list = [], [], []
                            for kp, val in knowledge_map.items():
                                my_rate = round((val['my']/val['full'])*100, 1) if val['full']>0 else 0
                                avg_rate = round((val['class_total']/val['full'])*100, 1) if val['full']>0 else 0
                                k_data.append({'知识点': kp, '我的掌握率': my_rate, '班级平均': avg_rate})
                                if my_rate < avg_rate: weak_points_list.append(kp)
                                else: strong_points_list.append(kp)
                            
                            df_kp = pd.DataFrame(k_data)
                            if not df_kp.empty:
                                c_chart, c_text = st.columns([1.2, 1])
                                with c_chart:
                                    fig = go.Figure()
                                    cats = df_kp['知识点'].tolist() + [df_kp['知识点'].tolist()[0]]
                                    mys = df_kp['我的掌握率'].tolist() + [df_kp['我的掌握率'].tolist()[0]]
                                    avgs = df_kp['班级平均'].tolist() + [df_kp['班级平均'].tolist()[0]]
                                    fig.add_trace(go.Scatterpolar(r=avgs, theta=cats, fill='toself', name='班级平均', line_color='#cccccc'))
                                    fig.add_trace(go.Scatterpolar(r=mys, theta=cats, fill='toself', name='我的掌握', line_color='#FF4B4B'))
                                    fig.update_layout(polar=dict(radialaxis=dict(visible=True, range=[0, 100])), paper_bgcolor='rgba(0,0,0,0)', dragmode=False)
                                    st.plotly_chart(fig, use_container_width=True, config=CHART_CONFIG)
                                with c_text:
                                    st.markdown("#### 🩺 专家系统诊断")
                                    if weak_points_list:
                                        for row in k_data:
                                            if row['知识点'] in weak_points_list:
                                                st.write(f"▪ **{row['知识点']}** (落后 {row['班级平均'] - row['我的掌握率']:.1f}%)")
                                    else: st.success("🎉 所有知识点均达标！")
                                
                                st.divider()
                                
                                # ================== 🔴 学生端 AI 分析与导出 ==================
                                ai_state_key = f"ai_stu_{st.session_state.logged_in_id}_{sel_sub}"
                                if AI_API_KEY:
                                    if st.button(f"✨ 提取专家 AI 提分建议", type="primary"):
                                        st.session_state[ai_state_key] = True

                                    if st.session_state.get(ai_state_key, False):
                                        with st.spinner("AI 导师正在云端调取档案..."):
                                            w_str = "、".join(weak_points_list) if weak_points_list else "无"
                                            s_str = "、".join(strong_points_list) if strong_points_list else "无"
                                            ai_reply = get_ai_advice_for_student(st.session_state.logged_in_student, sel_sub, w_str, s_str)
                                            
                                            st.markdown(f"<div class='ai-box'><b>AI导师：</b><br><br>{ai_reply}</div>", unsafe_allow_html=True)
                                            
                                            doc_title = f"{st.session_state.logged_in_student}_{sel_sub}_AI提分计划"
                                            file_data, file_name, mime_type = generate_ai_doc(doc_title, ai_reply)
                                            st.download_button(label="📥 一键保存 AI 分析报告 (支持 Word/TXT)", data=file_data, file_name=file_name, mime=mime_type)

# ==============================================================================
# 🚀 页面逻辑：教师后台 
# ==============================================================================
elif selected_nav == "教师后台":
    
    if not st.session_state.teacher_role:
        st.markdown("<h1 class='main-title'>🏫 教务与教师管理中枢</h1>", unsafe_allow_html=True)
        col_left, col_mid, col_right = st.columns([1, 1.8, 1])
        with col_left:
            st.markdown("<br><br>", unsafe_allow_html=True)
            if os.path.exists("panda.gif"): st.image("panda.gif", use_container_width=True)
            
        with col_mid:
            with st.container(border=True):
                st.markdown("<h3 style='text-align: center; color: #555;'>👨‍🏫 教职工专属通道</h3><br>", unsafe_allow_html=True)
                t_name = st.text_input("👤 您的姓名 (需与学校花名册一致)")
                pwd = st.text_input("🔐 专属密码", type="password")
                
                if st.button("验证并进入控制台", use_container_width=True, type="primary"):
                    try:
                        roster_df = pd.read_csv(URL_TEACHER_ROSTER, on_bad_lines='skip')
                    except:
                        roster_df = None
                        
                    if roster_df is not None and '教师姓名' in roster_df.columns:
                        t_info = roster_df[roster_df['教师姓名'].astype(str).str.strip() == t_name.strip()]
                        if not t_info.empty:
                            info = t_info.iloc[0]
                            actual_role = str(info.get('角色', '')).strip()
                            
                            is_auth = False
                            if actual_role == "教务处" and pwd == ADMIN_PASSWORD: is_auth = True
                            elif actual_role == "班主任" and (pwd == HOMEROOM_PASSWORD or pwd == ADMIN_PASSWORD): is_auth = True
                            elif actual_role == "任课教师" and (pwd == TEACHER_PASSWORD or pwd == ADMIN_PASSWORD): is_auth = True
                            
                            if is_auth:
                                st.session_state.teacher_role = actual_role
                                st.session_state.teacher_name = t_name.strip()
                                st.session_state.teacher_subject = str(info.get('学科', '')).strip()
                                
                                classes_raw = str(info.get('管理班级', ''))
                                classes_clean = re.sub(r'[，、。；/|\s]+', ',', classes_raw)
                                st.session_state.teacher_classes = [normalize_class_name(c) for c in classes_clean.split(',') if c.strip()]
                                st.rerun()
                            else: st.error("❌ 密码错误，请核对您所在级别的专属密码。")
                        else: st.error(f"❌ 权限表中未找到【{t_name}】的授权信息。")
                    else: st.error("⚠️ 无法读取教师权限表，请检查 URL_TEACHER_ROSTER。")
                    
        with col_right:
            st.markdown("<br><br>", unsafe_allow_html=True)
            if os.path.exists("star.gif"): st.image("star.gif", use_container_width=True)
            
    else:
        role = st.session_state.teacher_role
        name = st.session_state.teacher_name
        subject = st.session_state.teacher_subject
        my_classes = st.session_state.teacher_classes
        
        c1, c2 = st.columns([5, 1])
        c1.markdown(f"### 👨‍🏫 欢迎，{name}老师！【权限级别：{role}】")
        if c2.button("🚪 安全退出 (切换身份必点)", use_container_width=True): logout()
        
        master_df = build_master_df(LATEST_EXAM_IDX)
        if master_df is not None:
            adm_direction = st.selectbox("👉 选择分析群体", ["物理方向", "历史方向", "综合方向"])
            df_direction_global = master_df[master_df['方向'] == adm_direction]
            
            class_avg_global = pd.DataFrame()
            if not df_direction_global.empty and '总分' in df_direction_global.columns:
                class_avg_global = df_direction_global.groupby('班级')['总分'].mean().round(1).reset_index()
                class_avg_global['均分排名'] = class_avg_global['总分'].rank(ascending=False, method='min').astype(int)
            
            df_filtered = df_direction_global.copy()
            if role == "教务处": 
                st.success("👑 全局最高权限，可查看所有班级数据。")
            elif role in ["班主任", "任课教师"]:
                def class_match(cls_str):
                    c1 = normalize_class_name(cls_str)
                    for my_c in my_classes:
                        if my_c in c1 or c1 in my_c: return True
                    return False
                df_filtered = df_filtered[df_filtered['班级'].apply(class_match)]
                
                if role == "班主任": st.success(f"🛡️ 班级保护生效：仅查看【{'、'.join(my_classes)}】全科成绩。")
                else: st.success(f"🛡️ 单科保护生效：仅查看【{'、'.join(my_classes)}】的【{subject}】成绩。")
            
            if not df_filtered.empty:
                st.divider()
                
                if role == "班主任" and not class_avg_global.empty:
                    st.markdown(f"#### 🏆 【{adm_direction}】本班均分排名快报")
                    metric_cols = st.columns(len(my_classes))
                    for i, cls in enumerate(my_classes):
                        cls_info = class_avg_global[class_avg_global['班级'].apply(lambda x: normalize_class_name(cls) in normalize_class_name(x))]
                        if not cls_info.empty:
                            avg = cls_info.iloc[0]['总分']
                            rk = cls_info.iloc[0]['均分排名']
                            metric_cols[i].metric(label=f"{cls} 平均分", value=f"{avg} 分", delta=f"该方向年级第 {rk} 名", delta_color="off")
                        else:
                            metric_cols[i].metric(label=f"{cls}", value="暂无数据")
                    st.markdown("<br>", unsafe_allow_html=True)

                st.markdown(f"#### 📊 【{LATEST_EXAM['name']}】学情分析图表")

                # --- 🔴 任课教师视角 ---
                if role == "任课教师":
                    if subject in df_filtered.columns:
                        class_avg = df_filtered.groupby('班级')[subject].mean().round(1).reset_index()
                        fig_bar_t = px.bar(class_avg, x='班级', y=subject, text_auto=True, title=f"所带班级【{subject}】均分对比")
                        fig_bar_t.update_layout(dragmode=False)
                        st.plotly_chart(fig_bar_t, use_container_width=True, config=CHART_CONFIG)
                        
                        st.markdown("#### 📋 学生成绩明细表")
                        df_filtered[f'{subject}班级排名'] = df_filtered.groupby('班级')[subject].rank(ascending=False, method='min').fillna(0).astype(int)
                        df_filtered[f'{subject}年级排名'] = df_filtered[subject].rank(ascending=False, method='min').fillna(0).astype(int)

                        table_to_show = df_filtered[['姓名', '考号', '班级', f'{subject}年级排名', f'{subject}班级排名', subject]].sort_values(by=subject, ascending=False)
                        
                        render_html_table(table_to_show)
                        
                        excel_title = f"【{LATEST_EXAM['name']}】{subject}成绩单"
                        file_data, file_name, mime_type = generate_excel_download(table_to_show, f"{LATEST_EXAM['name']}_{subject}成绩单", excel_title)
                        st.download_button(label="📥 一键下载精美 Excel 成绩单", data=file_data, file_name=file_name, mime=mime_type, type="primary")
                        
                        if LATEST_EXAM.get(subject):
                            st.divider()
                            st.markdown(f"#### 🧠 【{subject}】底层共性诊断与 AI 教研")
                            df_diag = load_data(LATEST_EXAM[subject], header_lines=[0, 1, 2])
                            if df_diag is not None:
                                name_c, id_c, cls_c = None, None, None
                                for col in df_diag.columns:
                                    cstr = str(col[0]) if isinstance(col, tuple) else str(col)
                                    if '姓名' in cstr: name_c = col
                                    elif '考号' in cstr or '学号' in cstr: id_c = col
                                    elif '班级' in cstr: cls_c = col
                                    
                                k_stats = {}
                                for col in df_diag.columns:
                                    if col in [name_c, id_c, cls_c]: continue
                                    cstr = str(col[0]) if isinstance(col, tuple) else str(col)
                                    if '总分' in cstr or '排名' in cstr: continue
                                    
                                    q_name = str(col[0]).strip() if isinstance(col, tuple) else str(col).strip()
                                    kp = str(col[1]).strip() if isinstance(col, tuple) and len(col) > 1 else q_name
                                    if kp == "" or kp.startswith("Unnamed"): kp = q_name
                                    
                                    try: full = float(col[2]) if isinstance(col, tuple) and len(col) > 2 else 0
                                    except: full = 0
                                    
                                    if full <= 0:
                                        try: full = float(pd.to_numeric(df_diag[col], errors='coerce').max())
                                        except: full = 0
                                        
                                    if full > 0:
                                        if kp not in k_stats: k_stats[kp] = []
                                        k_stats[kp].append(pd.to_numeric(df_diag[col], errors='coerce').mean() / full)
                                        
                                if k_stats:
                                    k_final = [{"知识点": kp, "掌握率": round(sum(rates)/len(rates)*100, 1)} for kp, rates in k_stats.items()]
                                    df_k = pd.DataFrame(k_final).sort_values("掌握率")
                                    fig_k = px.bar(df_k, x="掌握率", y="知识点", orientation='h')
                                    fig_k.update_layout(dragmode=False)
                                    st.plotly_chart(fig_k, use_container_width=True, config=CHART_CONFIG)
                                    
                                    # ================== 🔴 教师端 AI 分析与导出 ==================
                                    ai_t_key = f"ai_tea_{subject}"
                                    if AI_API_KEY:
                                        if st.button("✨ 提取专家 AI 教研建议", type="primary"):
                                            st.session_state[ai_t_key] = True

                                        if st.session_state.get(ai_t_key, False):
                                            with st.spinner("AI 正在云端调取报告..."):
                                                ai_reply = get_ai_advice_for_teacher(subject, '、'.join(df_k.head(3)['知识点'].tolist()))
                                                st.markdown(f"<div class='ai-box'>{ai_reply}</div>", unsafe_allow_html=True)
                                                
                                                doc_title = f"【{LATEST_EXAM['name']}】{subject}学科_AI教研指导报告"
                                                file_data, file_name, mime_type = generate_ai_doc(doc_title, ai_reply)
                                                st.download_button(label="📥 一键保存 AI 教研报告 (支持 Word/TXT)", data=file_data, file_name=file_name, mime=mime_type)

                    else: st.warning(f"当前群体的考试中未找到您的学科【{subject}】。")
                
                # --- 教务处 & 班主任视角 ---
                else:
                    class_avg = df_filtered.groupby('班级')['总分'].mean().round(1).reset_index()
                    fig_bar_tot = px.bar(class_avg, x='班级', y='总分', text_auto=True, title="总分均分对照")
                    fig_bar_tot.update_layout(dragmode=False)
                    st.plotly_chart(fig_bar_tot, use_container_width=True, config=CHART_CONFIG)
                    
                    st.markdown("#### 📋 学生全科成绩明细表")
                    cols = df_filtered.columns.tolist()
                    front_cols = ['姓名', '考号', '班级', '总分', '总分年级排名', '总分班级排名', '方向']
                    other_cols = [c for c in cols if c not in front_cols]
                    table_to_show = df_filtered[front_cols + other_cols].sort_values(by=['班级', '总分'], ascending=[True, False])
                    
                    render_html_table(table_to_show)
                    
                    excel_title = f"【{LATEST_EXAM['name']}】{adm_direction}成绩汇总单"
                    file_data, file_name, mime_type = generate_excel_download(table_to_show, f"{LATEST_EXAM['name']}_{adm_direction}成绩汇总单", excel_title)
                    st.download_button(label="📥 一键下载精美 Excel 成绩单", data=file_data, file_name=file_name, mime=mime_type, type="primary")

                # ==========================================================
                # 🌟 终极神技：个人历次成绩档案智能搜索追踪系统！
                # ==========================================================
                st.divider()
                st.markdown("#### 🔍 个人历次成绩档案检索")
                st.info("💡 在下方下拉框中搜索或选择学生姓名，即可一键调取该生所有历史考试的成绩波动轨迹！")
                
                student_options = df_filtered.apply(lambda x: f"{x['班级']} | {x['姓名']} | 考号:{x['考号']}", axis=1).tolist()
                
                if student_options:
                    sel_student_str = st.selectbox("请搜索或选择目标学生：", ["-- 请选择学生 --"] + student_options)
                    
                    if sel_student_str != "-- 请选择学生 --":
                        sel_id = clean_str(sel_student_str.split("考号:")[1].strip())
                        sel_name = sel_student_str.split("|")[1].strip()
                        
                        history_records = []
                        with st.spinner(f"正在云端数据库检索【{sel_name}】的历次考试档案..."):
                            for i, exam in enumerate(EXAMS):
                                m_df = build_master_df(i)
                                if m_df is not None and not m_df.empty:
                                    stu_h = m_df[m_df['考号'] == sel_id]
                                    if not stu_h.empty:
                                        row = stu_h.iloc[0]
                                        rec = {"考试名称": exam['name']}
                                        
                                        if '总分' in row: rec['总分'] = float(row['总分'])
                                        if '总分年级排名' in row: rec['年级排名'] = int(row['总分年级排名'])
                                        
                                        for s in ['语文','数学','英语','物理','化学','生物','历史','政治','地理']:
                                            if s in row and pd.notna(row[s]):
                                                rec[s] = float(row[s])
                                        history_records.append(rec)
                                        
                        if history_records:
                            df_hist = pd.DataFrame(history_records)
                            
                            if role == "任课教师":
                                if subject in df_hist.columns:
                                    fig = px.line(df_hist, x="考试名称", y=subject, markers=True, title=f"【{sel_name}】{subject} 历次成绩走势", line_shape="spline")
                                    fig.update_traces(line_color="#FF4B4B", marker=dict(size=10))
                                    fig.update_layout(dragmode=False)
                                    st.plotly_chart(fig, use_container_width=True, config=CHART_CONFIG)
                                else:
                                    st.warning(f"未检索到该生【{subject}】的历史成绩。")
                            else:
                                t_col1, t_col2 = st.columns(2)
                                with t_col1:
                                    if "总分" in df_hist.columns:
                                        fig1 = px.line(df_hist, x="考试名称", y="总分", markers=True, title=f"【{sel_name}】历次总分走势", line_shape="spline")
                                        fig1.update_traces(line_color="#FF4B4B", marker=dict(size=10))
                                        fig1.update_layout(dragmode=False)
                                        st.plotly_chart(fig1, use_container_width=True, config=CHART_CONFIG)
                                with t_col2:
                                    if "年级排名" in df_hist.columns:
                                        fig2 = px.line(df_hist, x="考试名称", y="年级排名", markers=True, title=f"【{sel_name}】历次总分年级排名走势", line_shape="spline")
                                        fig2.update_traces(line_color="#0068C9", marker=dict(size=10))
                                        fig2.update_yaxes(autorange="reversed")
                                        fig2.update_layout(dragmode=False)
                                        st.plotly_chart(fig2, use_container_width=True, config=CHART_CONFIG)
                                
                                st.markdown("##### 🔬 该生单科历史透视")
                                avail_hist_subs = [s for s in ['语文','数学','英语','物理','化学','生物','历史','政治','地理'] if s in df_hist.columns]
                                if avail_hist_subs:
                                    sel_hist_sub = st.selectbox("选择要查看的单科：", avail_hist_subs, key="hist_sub")
                                    fig3 = px.line(df_hist, x="考试名称", y=sel_hist_sub, markers=True, title=f"【{sel_name}】{sel_hist_sub} 历次成绩走势", line_shape="spline")
                                    fig3.update_traces(line_color="#10B981", marker=dict(size=10))
                                    fig3.update_layout(dragmode=False)
                                    st.plotly_chart(fig3, use_container_width=True, config=CHART_CONFIG)
                        else:
                            st.info("暂未抓取到该生的历史轨迹。")
            else: st.warning("⚠️ 在当前选择的群体中，未找到您的授权班级数据。")