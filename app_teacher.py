import streamlit as st
import pandas as pd
import plotly.graph_objects as go
import plotly.express as px
from streamlit_option_menu import option_menu
import openai
import os
import functools
import io
import re

# ==============================================================================
# 1. 页面基础配置 (SaaS 宽屏模式)
# ==============================================================================
st.set_page_config(page_title="英华教务教研指挥舱", layout="wide", page_icon="🏢", initial_sidebar_state="expanded")

# ==============================================================================
# 🎨 顶级 SaaS 美学 CSS 样式注入 (极简、紧凑、去白边)
# ==============================================================================
st.markdown("""
<style>
    #MainMenu {visibility: hidden;} 
    footer {visibility: hidden;} 
    header {background: transparent !important;}
    
    .block-container { padding-top: 1rem !important; padding-bottom: 2rem !important; max-width: 96% !important;}
    .stApp { background-color: #F4F7F9; }
    
    [data-testid="stSidebar"] {
        min-width: 250px !important; max-width: 250px !important; background-color: #0B1120 !important;
    }
    [data-testid="stSidebar"] p, [data-testid="stSidebar"] label, [data-testid="stSidebar"] span { color: #94A3B8 !important; }
    [data-testid="stSidebar"] h1, [data-testid="stSidebar"] h2, [data-testid="stSidebar"] h3 { color: #F8FAFC !important; }
    [data-testid="stSidebar"] hr { border-color: rgba(255,255,255,0.05) !important; margin-top: 10px !important; margin-bottom: 10px !important; }

    [data-testid="stSidebar"] div[data-testid="stSelectbox"] { margin-bottom: -15px !important; }
    [data-testid="stSidebar"] div[data-baseweb="select"] > div {
        background-color: #1E293B !important; border: 1px solid rgba(255,255,255,0.05) !important; border-radius: 8px !important;
    }
    [data-testid="stSidebar"] div[data-baseweb="select"] div { color: #E2E8F0 !important; }
    [data-testid="stSidebar"] div[data-baseweb="select"] svg { fill: #94A3B8 !important; }
    ul[role="listbox"] { background-color: #1E293B !important; }
    li[role="option"] { background-color: #1E293B !important; color: #E2E8F0 !important; }
    li[role="option"]:hover { background-color: #334155 !important; }

    [data-testid="stSidebar"] [data-testid="stExpander"] { border: none !important; background: transparent !important; }
    [data-testid="stSidebar"] [data-testid="stExpander"] details, 
    [data-testid="stSidebar"] [data-testid="stExpander"] summary {
        background: transparent !important; border: none !important; padding: 0 !important; margin-bottom: 5px !important;
    }
    [data-testid="stSidebar"] [data-testid="stExpander"] summary p { font-size: 15px !important; font-weight: bold !important; color: #94A3B8 !important; }
    [data-testid="stSidebar"] [data-testid="stExpander"] summary svg { fill: #64748B !important; }
    [data-testid="stSidebar"] [data-testid="stExpander"] [data-testid="stExpanderDetails"] {
        background-color: #162032 !important; border-radius: 10px !important; padding: 0 !important; overflow: hidden !important; border: 1px solid rgba(255,255,255,0.03) !important;
    }
    [data-testid="stSidebar"] iframe { background-color: transparent !important; }

    [data-testid="stSidebar"] div.stButton > button {
        background-color: #1E293B !important; color: #94A3B8 !important; 
        border: 1px solid rgba(255,255,255,0.05) !important; border-radius: 8px !important;
        font-weight: bold !important; padding-top: 10px !important; padding-bottom: 10px !important; transition: all 0.3s ease;
    }
    [data-testid="stSidebar"] div.stButton > button:hover { background-color: #334155 !important; color: #F8FAFC !important; border-color: rgba(255,255,255,0.1) !important; }

    .header-card {
        background-color: #FFFFFF; padding: 15px 25px; border-radius: 10px;
        box-shadow: 0 2px 10px rgba(0,0,0,0.03); display: flex; justify-content: space-between;
        align-items: center; margin-bottom: 20px; border-left: 6px solid #3B82F6;
    }
    div[data-testid="stMetric"] { background-color: #FFFFFF; border-radius: 10px; padding: 15px; box-shadow: 0 2px 10px rgba(0,0,0,0.02); border: 1px solid #F1F5F9; text-align: center; }
    div[data-testid="stMetric"] label { font-size: 14px !important; color: #64748B !important; }
    div[data-testid="stMetric"] div[data-testid="stMetricValue"] { font-size: 26px !important; color: #0F172A !important; font-weight: 800 !important; }
    div[data-testid="stForm"] { background-color: #ffffff; padding: 40px; border-radius: 15px; box-shadow: 0 8px 20px rgba(0,0,0,0.05); border: none; }
    .ai-box { background: #FFFFFF; border-left: 5px solid #10B981; padding: 20px; border-radius: 8px; font-size: 15px; color: #333; line-height: 1.8; box-shadow: 0 2px 10px rgba(0,0,0,0.03);}
</style>
""", unsafe_allow_html=True)

CHART_CONFIG = {'displayModeBar': False, 'scrollZoom': False}

# ==============================================================================
# 🔐 权限组与环境读取
# ==============================================================================
GLOBAL_ROLES = ["校长", "副校长", "教学主任", "教务处"]
SUBJECT_HEAD_ROLES = ["学科主任"]
TEACHER_ROLES = ["任课教师", "教师"]
HOMEROOM_ROLES = ["班主任"]

if 'current_grade' not in st.session_state: st.session_state.current_grade = "高三"

try:
    ADMIN_PASSWORD = st.secrets["ADMIN_PWD"]
    HOMEROOM_PASSWORD = st.secrets.get("HOMEROOM_PWD", ADMIN_PASSWORD) 
    TEACHER_PASSWORD = st.secrets.get("TEACHER_PWD", ADMIN_PASSWORD)
    URL_TEACHER_ROSTER = st.secrets.get("URL_TEACHER_ROSTER", "") 
    URL_EXAM_CONFIG = st.secrets.get("URL_EXAM_CONFIG", "")
    AI_API_KEY = st.secrets.get("DEEPSEEK_API_KEY", "")
except Exception as e:
    st.error("⚠️ 系统配置读取失败，请检查 Streamlit 后台的 Secrets。")
    st.stop()

if AI_API_KEY: client = openai.OpenAI(api_key=AI_API_KEY, base_url="https://api.deepseek.com")
else: client = None

# ==============================================================================
# 🛠️ 核心引擎与数据缓存 (🛡️ 彻底修复无考号闪退)
# ==============================================================================
def clean_url(url): return str(url).strip() if pd.notna(url) and str(url).strip().lower() != 'nan' else ""
def clean_str(val): return str(val).strip()[:-2] if pd.notna(val) and str(val).strip().endswith('.0') else str(val).strip() if pd.notna(val) else ""
def clean_name(val): return str(val).replace(" ", "").strip() if pd.notna(val) else ""

@st.cache_data(ttl=300)
def load_exam_config(url):
    try: return pd.read_csv(url, on_bad_lines='skip')
    except: return pd.DataFrame()

def normalize_class_name(c):
    if pd.isna(c): return "未分班"
    c = str(c).replace(" ", "").strip()
    if not c: return "未分班"
    for k, v in {'1':'一','2':'二','3':'三','4':'四','5':'五','6':'六','7':'七','8':'八','9':'九','0':'零'}.items(): c = c.replace(k, v)
    c = c.replace("高三","").replace("高二","").replace("高一","").replace("年级","").replace("()","").replace("（）","")
    if not c.endswith("班"): c += "班"
    return c

@st.cache_data(ttl=600)
def load_data(url, header_lines=0):
    if not url or not url.strip(): return None
    try: return pd.read_csv(url, header=header_lines, on_bad_lines='skip')
    except: return None

@st.cache_data(ttl=600, show_spinner=False)
def build_master_df(grade_key):
    config_df = load_exam_config(URL_EXAM_CONFIG)
    exams = []
    if not config_df.empty and '年级' in config_df.columns:
        grade_config = config_df[config_df['年级'].astype(str).str.strip() == grade_key]
        for _, row in grade_config.iterrows():
            exams.append({
                "name": str(row.get('考试名称', '')).strip(),
                "语文": clean_url(row.get('语文')), "数学": clean_url(row.get('数学')), "英语": clean_url(row.get('英语')),
                "物理": clean_url(row.get('物理')), "化学": clean_url(row.get('化学')), "生物": clean_url(row.get('生物')),
                "历史": clean_url(row.get('历史')), "政治": clean_url(row.get('政治')), "地理": clean_url(row.get('地理'))
            })
    if not exams: return None, None, []
    
    latest_exam = exams[-1]
    dfs = []
    subs = ['语文','数学','英语','物理','化学','生物','历史','政治','地理']
    for sub in subs:
        url = latest_exam.get(sub)
        if url:
            df_sub = load_data(url, header_lines=[0,1,2])
            if df_sub is not None:
                name_c, id_c, cls_c = None, None, None
                for col in df_sub.columns:
                    cstr = str(col[0]) if isinstance(col, tuple) else str(col)
                    if '姓名' in cstr: name_c = col
                    elif '考号' in cstr or '学号' in cstr: id_c = col
                    elif '班级' in cstr: cls_c = col
                
                # 🛡️ 无考号盲合补丁：只要有姓名，哪怕没考号也强制合并！没考号就把名字当考号用！
                if name_c:
                    res = []
                    for _, row in df_sub.iterrows():
                        tot = sum(float(row[c]) for c in df_sub.columns if c not in [name_c, id_c, cls_c] and '总分' not in str(c) and '排名' not in str(c) and pd.notna(row[c]) and str(row[c]).replace('.','',1).isdigit())
                        s_name = clean_name(row[name_c])
                        # 如果没有id_c列，或者id_c列为空，直接把姓名当唯一标识符赋给考号，绝不报错
                        s_id = clean_str(row[id_c]) if id_c and pd.notna(row[id_c]) else s_name 
                        s_cls = normalize_class_name(row[cls_c]) if cls_c else "未分班"
                        
                        if s_name: res.append({'姓名': s_name, '考号': s_id, '班级': s_cls, sub: round(tot, 1)})
                    
                    if res: dfs.append(pd.DataFrame(res))
                    
    if not dfs: return None, latest_exam, exams
    
    # 强制合并，所有 dataframe 此时绝对都包含了 '姓名','考号','班级'
    master = functools.reduce(lambda l, r: pd.merge(l, r, on=['姓名','考号','班级'], how='outer'), dfs)
    present_subs = [s for s in subs if s in master.columns]
    master[present_subs] = master[present_subs].fillna(0)
    master['总分'] = master[present_subs].sum(axis=1).round(1)
    
    def get_dir(r):
        if r.get('物理', 0) > 0: return "物理方向"
        if r.get('历史', 0) > 0: return "历史方向"
        cls_n = str(r.get('班级', ''))
        if '文' in cls_n or '史' in cls_n: return "历史方向"
        if '理' in cls_n or '物' in cls_n: return "物理方向"
        return "综合方向"
    master['方向'] = master.apply(get_dir, axis=1)
    master['总分班级排名'] = master.groupby(['班级', '方向'])['总分'].rank(ascending=False, method='min').fillna(0).astype(int)
    master['总分年级排名'] = master.groupby(['方向'])['总分'].rank(ascending=False, method='min').fillna(0).astype(int)
    return master, latest_exam, exams

# ==============================================================================
# 🎨 导出与排版引擎 (含班级拆分)
# ==============================================================================
def render_html_table(df):
    html = """
    <div style="width: 100%; overflow-x: auto; margin-bottom: 25px; border-radius: 8px; box-shadow: 0 2px 8px rgba(0,0,0,0.03); border: 1px solid #E8E8E8; background: white;">
    <table style="width: 100%; border-collapse: collapse; font-size: 14px; text-align: center; font-family: 'Helvetica Neue', Arial, sans-serif;">
    """
    html += "<tr>" + "".join([f"<th style='background-color: #FAFAFA; color: #333; padding: 10px 12px; border-bottom: 2px solid #E8E8E8; white-space: nowrap; font-weight: 800;'>{col}</th>" for col in df.columns]) + "</tr>"
    for i, row in df.iterrows():
        bg_color = "#FFFFFF" if i % 2 == 0 else "#FAFAFA"
        html += f"<tr style='background-color: {bg_color}; transition: background-color 0.2s;' onmouseover=\"this.style.backgroundColor='#E6F7FF'\" onmouseout=\"this.style.backgroundColor='{bg_color}'\">"
        for col in df.columns:
            val = row[col]
            if isinstance(val, float): val = f"{val:.1f}"
            html += f"<td style='padding: 8px 12px; border-bottom: 1px solid #F0F0F0; color: #555;'>{val}</td>"
        html += "</tr>"
    html += "</table></div>"
    st.markdown(html, unsafe_allow_html=True)

def generate_excel_download(df, filename_prefix, title_text, split_by_class=False):
    try:
        from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
        from openpyxl.utils import get_column_letter
        buffer = io.BytesIO()
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            def write_sheet(dataframe, sheet_name, sheet_title):
                dataframe.to_excel(writer, index=False, sheet_name=sheet_name, startrow=1)
                worksheet = writer.sheets[sheet_name]
                num_cols = len(dataframe.columns)
                if num_cols == 0: return
                worksheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=num_cols)
                title_cell = worksheet.cell(row=1, column=1, value=sheet_title)
                title_cell.font = Font(size=18, bold=True, color="FFFFFFFF") 
                title_cell.fill = PatternFill(start_color="FF3B82F6", end_color="FF3B82F6", fill_type="solid") 
                title_cell.alignment = Alignment(horizontal="center", vertical="center")
                worksheet.row_dimensions[1].height = 40 
                header_fill = PatternFill(start_color="FF64748B", end_color="FF64748B", fill_type="solid")
                header_font = Font(bold=True, color="FFFFFFFF", size=11)
                even_fill = PatternFill(start_color="FFF8FAFC", end_color="FFF8FAFC", fill_type="solid")
                odd_fill = PatternFill(start_color="FFFFFFFF", end_color="FFFFFFFF", fill_type="solid")
                thin_border = Border(left=Side(style='thin', color='FFDDDDDD'), right=Side(style='thin', color='FFDDDDDD'), 
                                     top=Side(style='thin', color='FFDDDDDD'), bottom=Side(style='thin', color='FFDDDDDD'))
                for col_idx in range(1, num_cols + 1):
                    col_letter = get_column_letter(col_idx)
                    max_len = sum(2 if ord(c)>127 else 1 for c in str(dataframe.columns[col_idx-1]))
                    for row_idx in range(2, len(dataframe) + 3):
                        cell = worksheet.cell(row=row_idx, column=col_idx)
                        cell.alignment = Alignment(horizontal="center", vertical="center")
                        cell.border = thin_border
                        if row_idx == 2:
                            cell.fill = header_fill; cell.font = header_font
                        else:
                            cell.fill = even_fill if row_idx % 2 == 0 else odd_fill; cell.font = Font(size=11)
                        val_str = str(cell.value) if cell.value is not None else ""
                        val_len = sum(2 if ord(c)>127 else 1 for c in val_str)
                        if val_len > max_len: max_len = val_len
                    worksheet.column_dimensions[col_letter].width = max_len + 4
                for row_idx in range(2, len(dataframe) + 3): worksheet.row_dimensions[row_idx].height = 22

            write_sheet(df, '全年级汇总表', title_text)
            if split_by_class and '班级' in df.columns:
                classes = sorted(df['班级'].dropna().unique().tolist())
                for cls in classes:
                    cls_df = df[df['班级'] == cls]
                    safe_sheet_name = str(cls)[:30].replace('[','').replace(']','').replace('*','')
                    write_sheet(cls_df, safe_sheet_name, f"{title_text} - {cls}")
                    
        return buffer.getvalue(), f"{filename_prefix}.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    except Exception as e:
        return df.to_csv(index=False).encode('utf-8-sig'), f"{filename_prefix}.csv", "text/csv"

def generate_ai_doc(title, content):
    try:
        import docx
        from docx.shared import Pt
        from docx.oxml.ns import qn
        from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
        doc = docx.Document()
        style = doc.styles['Normal']; style.font.name = '仿宋'; style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋'); style.font.size = Pt(10.5)
        h = doc.add_heading(level=1); h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER; run = h.add_run(title)
        run.font.name = '黑体'; run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体'); run.font.size = Pt(14); run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
        def add_runs_to_paragraph(p, text):
            for part in re.split(r'(\*\*.*?\*\*)', text):
                r = p.add_run(part[2:-2].replace('*', '') if part.startswith('**') and part.endswith('**') else part.replace('*', ''))
                if part.startswith('**'): r.bold = True
                r.font.name = '仿宋'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋'); r.font.size = Pt(10.5)
        for line in content.split('\n'):
            line = line.strip()
            if not line: continue
            if line.startswith('- ') or line.startswith('* '): line = line[2:].strip()
            line = line.replace('•', '').replace('·', '').strip()
            p = doc.add_paragraph(); p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE 
            if line.startswith('#'):
                level = 0
                while level < len(line) and line[level] == '#': level += 1
                r = p.add_run(line[level:].strip().replace('*', '')); r.bold = True; r.font.size = Pt(12); r.font.name = '黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            else:
                add_runs_to_paragraph(p, line)
        buffer = io.BytesIO(); doc.save(buffer)
        return buffer.getvalue(), f"{title}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    except:
        return f"【{title}】\n\n{content}".encode('utf-8-sig'), f"{title}.txt", "text/plain"

# ==============================================================================
# 🧠 动态 AI 引擎
# ==============================================================================
@st.cache_data(ttl=2592000, show_spinner=False)
def get_ai_grouped_advice_for_teacher(grade, subject, grouped_data_str):
    if not client: return "⚠️ AI 尚未配置。"
    prompt = f"你是资深的{grade}{subject}教研专家。以下是我所带班级薄弱知识点及对应的具体学生名单（得分率不足60%）：\n{grouped_data_str}\n请生成一份「精准靶向辅导与分层教学报告」，深度剖析并给出具体措施，必须自然提及学生名字。"
    try: res = client.chat.completions.create(model="deepseek-chat", messages=[{"role": "system", "content": "你是精准教研专家AI。"}, {"role": "user", "content": prompt}]); return res.choices[0].message.content
    except: return "AI 生成失败"

@st.cache_data(ttl=2592000, show_spinner=False)
def get_ai_compare_advice(sheet_a, sheet_b, metric, class_avg_str, top_improvers_str, bottom_regressors_str):
    if not client: return "⚠️ AI 尚未配置。"
    
    if metric == "总分": focus_instruction = "【分析方向指令】：本次对比的是「总分」。请从各科均衡发展、时间分配、考试心态、整体复习策略等宏观教务管理角度，为年级主任和班主任提供策略。"
    else: focus_instruction = f"【分析方向指令】：本次对比的是单科「{metric}」。请深入该学科的知识体系，从微观教研、课堂教学改进、单科培优补差等专业学科角度，为{metric}备课组和任课教师提供落地方案。"

    prompt = f"""你是资深的高中教务数据分析专家。请基于以下真实考试数据，生成一份【{sheet_b}】对比【{sheet_a}】的【{metric}】学情进退步诊断报告。
{focus_instruction}
1. 各班【{metric}】平均进步幅度：
{class_avg_str}
2. 【{metric}】全校进步飞跃榜（表扬与经验总结）：
{top_improvers_str}
3. 【{metric}】需重点关注的退步生预警（归因与辅导建议）：
{bottom_regressors_str}
请使用专业、干练的语言，结构化输出报告，给出极具落地性的教研调整建议。"""
    try: res = client.chat.completions.create(model="deepseek-chat", messages=[{"role": "system", "content": "你是资深教务数据分析AI。"}, {"role": "user", "content": prompt}]); return res.choices[0].message.content
    except: return "AI 生成失败"

# ==============================================================================
# 🌟 智能解析器缓存提速 (🛡️ 加入空表头修复)
# ==============================================================================
@st.cache_data(ttl=600, show_spinner=False)
def smart_parse_excel(file_bytes, sheet_name):
    df_test = pd.read_excel(io.BytesIO(file_bytes), sheet_name=sheet_name, header=None, nrows=5)
    is_single_header = False
    if len(df_test) > 1:
        row_1_values = df_test.iloc[1].values
        if any(isinstance(x, (int, float)) and pd.notna(x) for x in row_1_values):
            is_single_header = True
            
    if is_single_header:
        df = pd.read_excel(io.BytesIO(file_bytes), sheet_name=sheet_name, header=0)
        cols = list(df.columns)
        for i in range(len(cols)):
            c_str = str(cols[i]).strip()
            if 'Unnamed' in c_str or c_str == '':
                if i == 1: cols[i] = '姓名'
                elif i == 0: cols[i] = '序号'
        df.columns = [str(c).strip() for c in cols]
        return df
    else:
        df = pd.read_excel(io.BytesIO(file_bytes), sheet_name=sheet_name, header=[0, 1])
        new_cols = []
        for c in df.columns:
            c0, c1 = str(c[0]).strip(), str(c[1]).strip()
            if 'Unnamed' in c0: c0 = ''
            if 'Unnamed' in c1: c1 = ''
            if c0 and c1 and c0 != c1: new_cols.append(f"{c0}_{c1}")
            elif c0: new_cols.append(c0)
            elif c1: new_cols.append(c1)
            else: new_cols.append("未知列")
        df.columns = new_cols
        return df

def find_target_column(df_columns, keywords):
    for kw in keywords:
        for col in df_columns:
            if kw in col: return col
    return None

# ==============================================================================
# 🛡️ 状态管理
# ==============================================================================
if 'teacher_role' not in st.session_state: st.session_state.teacher_role = None 
if 'teacher_name' not in st.session_state: st.session_state.teacher_name = None
if 'teacher_subject' not in st.session_state: st.session_state.teacher_subject = None
if 'teacher_classes' not in st.session_state: st.session_state.teacher_classes = []

def logout():
    for key in ['teacher_role', 'teacher_name', 'teacher_subject', 'teacher_classes']: st.session_state[key] = None
    st.rerun()

# ==============================================================================
# 🌐 左侧 SaaS 导航边栏
# ==============================================================================
menu_sel = "首页" 
adm_direction = "物理方向" 

if st.session_state.teacher_role:
    with st.sidebar:
        st.markdown(f"<h2 style='margin-top:-20px; padding-bottom: 10px;'>🏫 英华教务系统</h2>", unsafe_allow_html=True)
        st.markdown(f"""
        <div style='background: linear-gradient(135deg, #1E293B 0%, #0F172A 100%); padding: 18px; border-radius: 12px; margin-bottom: 20px; border: 1px solid rgba(255,255,255,0.05); box-shadow: 0 4px 12px rgba(0,0,0,0.2);'>
            <div style='color: #F8FAFC; font-size: 18px; font-weight: 800; letter-spacing: 1px;'>👨‍🏫 {st.session_state.teacher_name}</div>
            <div style='color: #94A3B8; font-size: 13px; margin-top: 6px;'>🛡️ 权限：{st.session_state.teacher_role}</div>
        </div>
        """, unsafe_allow_html=True)
        
        selected_grade = st.selectbox("隐藏标签1", ["高三", "高二", "高一"], index=["高三", "高二", "高一"].index(st.session_state.current_grade), label_visibility="collapsed")
        adm_direction = st.selectbox("隐藏标签2", ["物理方向", "历史方向", "综合方向"], label_visibility="collapsed")
        st.markdown("<div style='margin-top: 20px;'></div>", unsafe_allow_html=True)
        
        base_options = ["首页", "成绩明细表", "历次追踪分析", "AI 教研中心"]
        base_icons = ["grid-fill", "table", "graph-up-arrow", "robot"]
        if st.session_state.teacher_role in GLOBAL_ROLES:
            base_options.append("本地多考对比")
            base_icons.append("file-earmark-spreadsheet")
        
        with st.expander("📊 考试分析", expanded=True):
            menu_sel = option_menu(
                menu_title=None, options=base_options, icons=base_icons, menu_icon="cast", default_index=0,
                styles={
                    "container": {"padding": "5px 0!important", "background-color": "#162032", "border-radius": "0px", "border": "none"},
                    "icon": {"color": "#64748B", "font-size": "15px"},
                    "nav-link": {"font-size": "14px", "text-align": "left", "margin":"2px 10px", "color": "#64748B", "border-radius": "6px", "padding": "8px 15px"},
                    "nav-link-selected": {"background-color": "transparent", "color": "#F8FAFC", "font-weight": "bold"},
                    "nav-link:hover": {"background-color": "rgba(255,255,255,0.03)"}
                }
            )
        st.divider()
        st.button("🚪 退出", on_click=logout, use_container_width=True, type="secondary")
        if selected_grade != st.session_state.current_grade:
            st.session_state.current_grade = selected_grade; st.rerun()

# ==============================================================================
# 🚪 登录界面
# ==============================================================================
if not st.session_state.teacher_role:
    st.markdown("<br><br><br>", unsafe_allow_html=True)
    c_left, c_mid, c_right = st.columns([1, 1.5, 1])
    with c_mid:
        with st.form("teacher_login"):
            st.markdown(f"<h2 style='text-align: center; color: #0F172A; margin-bottom: 30px; font-weight: 800;'>🏫 英华数据指挥舱</h2>", unsafe_allow_html=True)
            t_name = st.text_input("👤 教职工姓名 (需与学校花名册一致)")
            pwd = st.text_input("🔐 访问密码", type="password")
            if st.form_submit_button("安全验证并进入", use_container_width=True):
                try: roster_df = pd.read_csv(URL_TEACHER_ROSTER, on_bad_lines='skip')
                except: roster_df = None
                if roster_df is not None and '教师姓名' in roster_df.columns:
                    t_info = roster_df[roster_df['教师姓名'].astype(str).str.strip() == t_name.strip()]
                    if not t_info.empty:
                        info = t_info.iloc[0]
                        actual_role = str(info.get('角色', '')).strip()
                        is_auth = False
                        if actual_role in GLOBAL_ROLES and pwd == ADMIN_PASSWORD: is_auth = True
                        elif actual_role in HOMEROOM_ROLES and (pwd == HOMEROOM_PASSWORD or pwd == ADMIN_PASSWORD): is_auth = True
                        elif actual_role in SUBJECT_HEAD_ROLES and (pwd == TEACHER_PASSWORD or pwd == ADMIN_PASSWORD): is_auth = True
                        elif actual_role in TEACHER_ROLES and (pwd == TEACHER_PASSWORD or pwd == ADMIN_PASSWORD): is_auth = True
                        if is_auth:
                            st.session_state.teacher_role = actual_role; st.session_state.teacher_name = t_name.strip()
                            st.session_state.teacher_subject = str(info.get('学科', '')).strip()
                            st.session_state.teacher_classes = [normalize_class_name(c) for c in re.sub(r'[，、。；/|\s]+', ',', str(info.get('管理班级', ''))).split(',') if c.strip()]
                            st.rerun()
                        else: st.error("❌ 密码错误或权限不匹配。")
                    else: st.error(f"❌ 权限表中未找到【{t_name}】。")
                else: st.error("⚠️ 无法读取教师权限表。")

# ==============================================================================
# 📊 核心业务区
# ==============================================================================
else:
    role = st.session_state.teacher_role
    name = st.session_state.teacher_name
    subject = st.session_state.teacher_subject
    my_classes = st.session_state.teacher_classes
    
    st.markdown(f"""
    <div class="header-card">
        <h3 style="margin: 0; color: #0F172A; font-weight: 800;">❖ {menu_sel} <span style="font-size:16px; color:#94A3B8; font-weight:normal; margin-left: 10px;">/ {st.session_state.current_grade} · {adm_direction}</span></h3>
    </div>
    """, unsafe_allow_html=True)
    
    # =========================================================================
    # 👑 本地多考对比分析
    # =========================================================================
    if menu_sel == "本地多考对比":
        st.info("💡 请上传包含各次考试独立Sheet的Excel文件。系统支持单层普通表头和双层复杂表头自动解析。")
        uploaded_file = st.file_uploader("📥 上传判卷系统 Excel 成绩表 (.xlsx / .xls)", type=["xlsx", "xls"])
        
        if uploaded_file:
            excel_data = pd.ExcelFile(uploaded_file)
            sheet_names = excel_data.sheet_names
            file_bytes = uploaded_file.getvalue() 
            
            with st.container(border=True):
                st.markdown("<p style='font-size: 15px; font-weight: bold; color: #0F172A;'>⚙️ 跨考比对引擎设置</p>", unsafe_allow_html=True)
                col1, col2, col3 = st.columns(3)
                sheet_a = col1.selectbox("基础考试 (A)", sheet_names)
                sheet_b = col2.selectbox("对比考试 (B)", sheet_names, index=min(1, len(sheet_names)-1))
                compare_metric = col3.selectbox("🎯 对比指标 (总分或单科)", ["总分", "语文", "数学", "英语", "物理", "化学", "生物", "历史", "政治", "地理"])
                
                df_a = smart_parse_excel(file_bytes, sheet_a)
                df_b = smart_parse_excel(file_bytes, sheet_b)
                
                name_col_a = find_target_column(df_a.columns, ['姓名', '名字'])
                name_col_b = find_target_column(df_b.columns, ['姓名', '名字'])
                
                if compare_metric == "总分":
                    metric_kws = ['总分赋分_分数', '总分_分数', '总分', '成绩', '赋分']
                    rank_kws = ['总分赋分_校名', '总分赋分_排名', '总分_校排', '总分_名次', '总分_校名', '总分_排名', '校名', '排名', '名次', '校排']
                else:
                    metric_kws = [f'{compare_metric}_分数', f'{compare_metric}_赋分', compare_metric]
                    rank_kws = [f'{compare_metric}_校名', f'{compare_metric}_排名', f'{compare_metric}_校排', f'{compare_metric}_名次']
                    
                score_col_a = find_target_column(df_a.columns, metric_kws)
                score_col_b = find_target_column(df_b.columns, metric_kws)
                rank_col_a = find_target_column(df_a.columns, rank_kws)
                rank_col_b = find_target_column(df_b.columns, rank_kws)
                class_col_b = find_target_column(df_b.columns, ['班级', '行政班', '总分赋分_班名', '总分_班名', '班名'])
                
                if not all([name_col_a, name_col_b, score_col_a, score_col_b]):
                    st.error(f"❌ 无法从表格中自动识别到「姓名」或「{compare_metric}」列，请检查 Excel 格式是否标准。")
                else:
                    base_score_col = f"{sheet_a}_{compare_metric}"
                    comp_score_col = f"{sheet_b}_{compare_metric}"
                    base_rank_col = f"{sheet_a}_校排名"
                    comp_rank_col = f"{sheet_b}_校排名"
                    
                    cols_a = {name_col_a: '姓名', score_col_a: base_score_col}
                    if rank_col_a: cols_a[rank_col_a] = base_rank_col
                    clean_a = df_a[list(cols_a.keys())].rename(columns=cols_a)
                    
                    cols_b = {name_col_b: '姓名', score_col_b: comp_score_col}
                    if rank_col_b: cols_b[rank_col_b] = comp_rank_col
                    if class_col_b: cols_b[class_col_b] = '班级'
                    clean_b = df_b[list(cols_b.keys())].rename(columns=cols_b)
                    
                    clean_a[base_score_col] = pd.to_numeric(clean_a[base_score_col], errors='coerce').fillna(0)
                    clean_b[comp_score_col] = pd.to_numeric(clean_b[comp_score_col], errors='coerce').fillna(0)
                    merged_df = pd.merge(clean_b, clean_a, on='姓名', how='inner')
                    
                    if merged_df.empty:
                        st.warning("⚠️ 两次考试中未能匹配到相同姓名的学生数据。请检查是否有同名不同字的情况。")
                    else:
                        merged_df['分数进步'] = (merged_df[comp_score_col] - merged_df[base_score_col]).round(1)
                        has_rank = False
                        if base_rank_col in merged_df.columns and comp_rank_col in merged_df.columns:
                            has_rank = True
                            merged_df[base_rank_col] = pd.to_numeric(merged_df[base_rank_col], errors='coerce')
                            merged_df[comp_rank_col] = pd.to_numeric(merged_df[comp_rank_col], errors='coerce')
                            merged_df['排名进步'] = merged_df[base_rank_col] - merged_df[comp_rank_col]
                            merged_df['排名进步'] = merged_df['排名进步'].fillna(0).astype(int)
                        
                        disp_cols = ['班级', '姓名', base_score_col, comp_score_col, '分数进步']
                        if has_rank: 
                            disp_cols.insert(3, base_rank_col)
                            disp_cols.insert(5, comp_rank_col)
                            disp_cols.append('排名进步')
                        disp_df = merged_df[disp_cols].sort_values('分数进步', ascending=False)
                        
                        st.markdown("---")
                        st.markdown(f"#### 🎯 【{sheet_b}】较【{sheet_a}】【{compare_metric}】全景分析")
                        
                        top_improvers = merged_df.sort_values(by='分数进步', ascending=False).head(5)
                        t_names = "、".join([f"【{r.get('班级', '未知')}】{r['姓名']}(+{r['分数进步']}分)" for _, r in top_improvers.iterrows()])
                        st.success(f"🏅 **全校进步飞跃榜 (Top 5)：** {t_names}")
                        
                        bottom_regressors = merged_df.sort_values(by='分数进步', ascending=True).head(5)
                        b_names = "、".join([f"【{r.get('班级', '未知')}】{r['姓名']}({r['分数进步']}分)" for _, r in bottom_regressors.iterrows()])
                        st.error(f"🚨 **需重点关注的退步预警 (Bottom 5)：** {b_names}")
                        
                        if '班级' in merged_df.columns:
                            class_avg_progress = merged_df.groupby('班级')['分数进步'].mean().round(1).reset_index()
                            class_avg_progress['Color'] = class_avg_progress['分数进步'].apply(lambda x: '#EF4444' if x < 0 else '#10B981')
                            fig_class = px.bar(class_avg_progress, x='班级', y='分数进步', text_auto=True, title="📊 各班级平均进步幅度横向对比")
                            fig_class.update_traces(marker_color=class_avg_progress['Color'], textposition='outside')
                            fig_class.update_layout(dragmode=False, plot_bgcolor='rgba(248, 250, 252, 0.5)', paper_bgcolor='white', margin=dict(t=40, b=20, l=20, r=20))
                            st.plotly_chart(fig_class, use_container_width=True, config=CHART_CONFIG)
                        
                        st.markdown("<br>", unsafe_allow_html=True)
                        c_table, c_chart = st.columns([1, 1])
                        with c_table:
                            st.markdown("<p style='font-size:14px; color:#64748B;'>📝 学生比对明细 (黄色背景为进步生)</p>", unsafe_allow_html=True)
                            def highlight_progress_yellow(row): return ['background-color: #FEF08A'] * len(row) if row.get('分数进步', 0) > 0 else [''] * len(row)
                            st.dataframe(disp_df.style.apply(highlight_progress_yellow, axis=1), hide_index=True, height=450, use_container_width=True)
                            
                            excel_title = f"【{sheet_b}】较【{sheet_a}】{compare_metric}进退步追踪表"
                            file_data, file_name, mime_type = generate_excel_download(disp_df, f"进退步对比_{compare_metric}", excel_title, split_by_class=True)
                            st.download_button(label="📥 下载精美比对结果 (已自动拆分班级Sheet)", data=file_data, file_name=file_name, mime=mime_type, type="primary", use_container_width=True)
                        
                        with c_chart:
                            st.markdown(f"<p style='font-size:14px; color:#64748B;'>📈 {compare_metric}四象限分布图 (红虚线为原地踏步线)</p>", unsafe_allow_html=True)
                            hover_d = ["班级", "分数进步"] if '班级' in merged_df.columns else ["分数进步"]
                            if has_rank: hover_d.append("排名进步")
                            fig = px.scatter(merged_df, x=base_score_col, y=comp_score_col, hover_name="姓名", hover_data=hover_d, color="分数进步", color_continuous_scale=px.colors.diverging.RdYlGn)
                            min_val = min(merged_df[base_score_col].min(), merged_df[comp_score_col].min()) - 10
                            max_val = max(merged_df[base_score_col].max(), merged_df[comp_score_col].max()) + 10
                            fig.add_shape(type="line", x0=min_val, y0=min_val, x1=max_val, y1=max_val, line=dict(color="#EF4444", dash="dash"))
                            fig.update_layout(dragmode=False, plot_bgcolor='rgba(248, 250, 252, 0.5)', paper_bgcolor='white', margin=dict(t=20, b=20, l=20, r=20), height=450, xaxis_title=f"X: {sheet_a} {compare_metric}", yaxis_title=f"Y: {sheet_b} {compare_metric}")
                            st.plotly_chart(fig, use_container_width=True, config=CHART_CONFIG)

                        # ==========================================
                        # 🎯 双考上线率进退步达标监控
                        # ==========================================
                        st.divider()
                        st.markdown(f"#### 🎯 【{compare_metric}】双考上线率进退步监控")
                        with st.container(border=True):
                            c_l1, c_l2 = st.columns(2)
                            with c_l1:
                                st.markdown(f"<p style='font-size:14px; font-weight:bold; color:#334155; margin-bottom:-5px;'>📌 设定【{sheet_a}】分数线</p>", unsafe_allow_html=True)
                                c_a_te, c_a_ben = st.columns(2)
                                line_te_a = c_a_te.number_input(f"🏆 {sheet_a} 特控线", min_value=0.0, value=0.0, step=1.0, key="te_a")
                                line_ben_a = c_a_ben.number_input(f"🎓 {sheet_a} 本科线", min_value=0.0, value=0.0, step=1.0, key="ben_a")
                            with c_l2:
                                st.markdown(f"<p style='font-size:14px; font-weight:bold; color:#334155; margin-bottom:-5px;'>📌 设定【{sheet_b}】分数线</p>", unsafe_allow_html=True)
                                c_b_te, c_b_ben = st.columns(2)
                                line_te_b = c_b_te.number_input(f"🏆 {sheet_b} 特控线", min_value=0.0, value=0.0, step=1.0, key="te_b")
                                line_ben_b = c_b_ben.number_input(f"🎓 {sheet_b} 本科线", min_value=0.0, value=0.0, step=1.0, key="ben_b")

                            if (line_te_a > 0 or line_ben_a > 0) or (line_te_b > 0 or line_ben_b > 0):
                                total_stu_comp = len(merged_df)
                                
                                base_te_all = len(merged_df[merged_df[base_score_col] >= line_te_a]) if line_te_a > 0 else 0
                                comp_te_all = len(merged_df[merged_df[comp_score_col] >= line_te_b]) if line_te_b > 0 else 0
                                base_ben_all = len(merged_df[merged_df[base_score_col] >= line_ben_a]) if line_ben_a > 0 else 0
                                comp_ben_all = len(merged_df[merged_df[comp_score_col] >= line_ben_b]) if line_ben_b > 0 else 0

                                rate_comp_te = round(comp_te_all / total_stu_comp * 100, 1) if total_stu_comp > 0 else 0
                                rate_comp_ben = round(comp_ben_all / total_stu_comp * 100, 1) if total_stu_comp > 0 else 0
                                
                                diff_te_count = comp_te_all - base_te_all
                                diff_ben_count = comp_ben_all - base_ben_all

                                st.markdown(f"""
                                <div style='background-color: #F8FAFC; padding: 15px; border-radius: 8px; border-left: 4px solid #10B981; margin-top: 15px; margin-bottom: 20px;'>
                                    <span style='font-size: 15px; color: #334155;'><b>🌐 双考群体达标波动概况：</b></span><br>
                                    <span style='font-size: 14px; color: #64748B;'>
                                    🏆 <b>{sheet_b}特控上线：</b> {comp_te_all} 人 (较{sheet_a} {'+' if diff_te_count>=0 else ''}{diff_te_count} 人)，当前特控率 <b style='color:#10B981;'>{rate_comp_te}%</b> &nbsp;&nbsp;|&nbsp;&nbsp; 
                                    🎓 <b>{sheet_b}本科上线：</b> {comp_ben_all} 人 (较{sheet_a} {'+' if diff_ben_count>=0 else ''}{diff_ben_count} 人)，当前本科率 <b style='color:#3B82F6;'>{rate_comp_ben}%</b>
                                    </span>
                                </div>
                                """, unsafe_allow_html=True)

                                if '班级' in merged_df.columns:
                                    class_stats = []
                                    for cls in sorted(merged_df['班级'].dropna().unique()):
                                        cdf = merged_df[merged_df['班级'] == cls]
                                        c_total = len(cdf)
                                        row_dict = {'班级': cls, '参考人数': c_total}
                                        
                                        if line_te_a > 0 or line_te_b > 0:
                                            c_base_te = len(cdf[cdf[base_score_col] >= line_te_a]) if line_te_a > 0 else 0
                                            c_comp_te = len(cdf[cdf[comp_score_col] >= line_te_b]) if line_te_b > 0 else 0
                                            row_dict[f'{sheet_a}特控数'] = c_base_te
                                            row_dict[f'{sheet_b}特控数'] = c_comp_te
                                            row_dict['特控上线变化'] = c_comp_te - c_base_te
                                            row_dict[f'{sheet_b}特控率(%)'] = round(c_comp_te / c_total * 100, 1) if c_total > 0 else 0
                                            
                                        if line_ben_a > 0 or line_ben_b > 0:
                                            c_base_ben = len(cdf[cdf[base_score_col] >= line_ben_a]) if line_ben_a > 0 else 0
                                            c_comp_ben = len(cdf[cdf[comp_score_col] >= line_ben_b]) if line_ben_b > 0 else 0
                                            row_dict[f'{sheet_a}本科数'] = c_base_ben
                                            row_dict[f'{sheet_b}本科数'] = c_comp_ben
                                            row_dict['本科上线变化'] = c_comp_ben - c_base_ben
                                            row_dict[f'{sheet_b}本科率(%)'] = round(c_comp_ben / c_total * 100, 1) if c_total > 0 else 0
                                            
                                        class_stats.append(row_dict)

                                    df_stats = pd.DataFrame(class_stats)
                                    if line_te_b > 0: df_stats = df_stats.sort_values(f'{sheet_b}特控率(%)', ascending=False)
                                    elif line_ben_b > 0: df_stats = df_stats.sort_values(f'{sheet_b}本科率(%)', ascending=False)

                                    html_stats = "<div style='width: 100%; overflow-x: auto; border-radius: 8px; border: 1px solid #E8E8E8;'><table style='width: 100%; border-collapse: collapse; font-size: 14px; text-align: center;'>"
                                    html_stats += "<tr><th style='background-color: #FAFAFA; padding: 10px; border-bottom: 2px solid #E8E8E8; color: #333;'>班级</th><th style='background-color: #FAFAFA; padding: 10px; border-bottom: 2px solid #E8E8E8; color: #333;'>参考人数</th>"
                                    if line_te_a > 0 or line_te_b > 0: html_stats += f"<th style='background-color: #FFFBEB; padding: 10px; border-bottom: 2px solid #E8E8E8; color: #92400E;'>{sheet_a}特控数</th><th style='background-color: #FFFBEB; padding: 10px; border-bottom: 2px solid #E8E8E8; color: #92400E;'>{sheet_b}特控数</th><th style='background-color: #FFFBEB; padding: 10px; border-bottom: 2px solid #E8E8E8; color: #92400E;'>特控变化</th><th style='background-color: #FFFBEB; padding: 10px; border-bottom: 2px solid #E8E8E8; color: #92400E;'>{sheet_b}特控率</th>"
                                    if line_ben_a > 0 or line_ben_b > 0: html_stats += f"<th style='background-color: #F0F9FF; padding: 10px; border-bottom: 2px solid #E8E8E8; color: #0369A1;'>{sheet_a}本科数</th><th style='background-color: #F0F9FF; padding: 10px; border-bottom: 2px solid #E8E8E8; color: #0369A1;'>{sheet_b}本科数</th><th style='background-color: #F0F9FF; padding: 10px; border-bottom: 2px solid #E8E8E8; color: #0369A1;'>本科变化</th><th style='background-color: #F0F9FF; padding: 10px; border-bottom: 2px solid #E8E8E8; color: #0369A1;'>{sheet_b}本科率</th>"
                                    html_stats += "</tr>"

                                    for _, r in df_stats.iterrows():
                                        html_stats += f"<tr><td style='padding: 8px; border-bottom: 1px solid #F0F0F0;'><b>{r['班级']}</b></td><td style='padding: 8px; border-bottom: 1px solid #F0F0F0;'>{r['参考人数']}</td>"
                                        if line_te_a > 0 or line_te_b > 0: 
                                            te_diff = r['特控上线变化']
                                            te_color = "red" if te_diff < 0 else "green"
                                            te_sign = "+" if te_diff > 0 else ""
                                            html_stats += f"<td style='padding: 8px; border-bottom: 1px solid #F0F0F0;'>{r[f'{sheet_a}特控数']}</td><td style='padding: 8px; border-bottom: 1px solid #F0F0F0; font-weight:bold;'>{r[f'{sheet_b}特控数']}</td><td style='padding: 8px; border-bottom: 1px solid #F0F0F0; color:{te_color};'>{te_sign}{te_diff}</td><td style='padding: 8px; border-bottom: 1px solid #F0F0F0; color:#B45309;'>{r[f'{sheet_b}特控率(%)']}%</td>"
                                        if line_ben_a > 0 or line_ben_b > 0: 
                                            ben_diff = r['本科上线变化']
                                            ben_color = "red" if ben_diff < 0 else "green"
                                            ben_sign = "+" if ben_diff > 0 else ""
                                            html_stats += f"<td style='padding: 8px; border-bottom: 1px solid #F0F0F0;'>{r[f'{sheet_a}本科数']}</td><td style='padding: 8px; border-bottom: 1px solid #F0F0F0; font-weight:bold;'>{r[f'{sheet_b}本科数']}</td><td style='padding: 8px; border-bottom: 1px solid #F0F0F0; color:{ben_color};'>{ben_sign}{ben_diff}</td><td style='padding: 8px; border-bottom: 1px solid #F0F0F0; color:#0284C7;'>{r[f'{sheet_b}本科率(%)']}%</td>"
                                        html_stats += "</tr>"
                                    html_stats += "</table></div>"

                                    st.markdown(html_stats, unsafe_allow_html=True)
                                    excel_title = f"【{sheet_b}】较【{sheet_a}】{compare_metric}上线率双考对比表"
                                    file_data, file_name, mime_type = generate_excel_download(df_stats, f"双考上线率对比_{compare_metric}", excel_title)
                                    st.download_button(label="📥 下载各班双考上线率对比报表 (Excel)", data=file_data, file_name=file_name, mime=mime_type, type="secondary")
                            else:
                                st.info(f"👆 请在上方输入框内设定两次考试的「特控线」或「本科线」，系统将立即为您计算班级达标率及边缘生转化情况！")

                        # ==========================================
                        # 🤖 AI 智能分析 (不会再闪退！)
                        # ==========================================
                        st.divider()
                        st.markdown("#### 🧠 全校学情进退步 AI 分析决策")
                        
                        # 使用独特的 key 将报告存在 session state 里，这样按钮点击也不会丢失
                        ai_compare_key = f"ai_comp_{sheet_a}_{sheet_b}_{compare_metric}"
                        
                        if st.button("✨ 一键生成全校质量分析报告 (基于底层数据的深层洞察)"):
                            with st.spinner("AI 正在汇聚全校数据，深度构建教学建议报告..."):
                                class_avg_str = "\n".join([f"- 【{r['班级']}】平均波动: {r['分数进步']}分" for _, r in class_avg_progress.iterrows()]) if '班级' in merged_df.columns else "未分班"
                                ai_reply = get_ai_compare_advice(sheet_a, sheet_b, compare_metric, class_avg_str, t_names, b_names)
                                st.session_state[ai_compare_key] = ai_reply

                        if ai_compare_key in st.session_state:
                            saved_reply = st.session_state[ai_compare_key]
                            st.markdown(f"<div class='ai-box'><b>🤖 专家指导建议：</b><br><br>{saved_reply}</div><br>", unsafe_allow_html=True)
                            
                            t_c1, t_c2, t_c3 = st.columns([1.5, 1, 1])
                            with t_c1:
                                if st.button("📌 存入本设备暂存库", key=f"save_btn_{ai_compare_key}"):
                                    saved_t_list_key = f"saved_ai_tea_list_comp"
                                    if saved_t_list_key not in st.session_state: st.session_state[saved_t_list_key] = []
                                    st.session_state[saved_t_list_key].insert(0, saved_reply)
                                    st.toast("✅ 已成功存入！")
                            with t_c3:
                                doc_title = f"【{sheet_b}】较【{sheet_a}】{compare_metric}_全校质量分析报告"
                                file_data, file_name, mime_type = generate_ai_doc(doc_title, saved_reply)
                                st.download_button(label="📥 下载 Word 格式分析报告", data=file_data, file_name=file_name, mime=mime_type, type="primary", key=f"down_btn_{ai_compare_key}")

    # =========================================================================
    # 常规模块
    # =========================================================================
    elif menu_sel != "本地多考对比":
        master_df, LATEST_EXAM, EXAMS_LIST = build_master_df(st.session_state.current_grade)
        
        if master_df is not None and not master_df.empty:
            df_direction_global = master_df[master_df['方向'] == adm_direction]
            class_avg_global = pd.DataFrame()
            if not df_direction_global.empty and '总分' in df_direction_global.columns:
                class_avg_global = df_direction_global.groupby('班级')['总分'].mean().round(1).reset_index()
                class_avg_global['均分排名'] = class_avg_global['总分'].rank(ascending=False, method='min').astype(int)
            
            df_filtered = df_direction_global.copy()
            if role in GLOBAL_ROLES: pass
            elif role in SUBJECT_HEAD_ROLES: pass 
            elif role in HOMEROOM_ROLES or role in TEACHER_ROLES:
                def class_match(cls_str):
                    c1 = normalize_class_name(cls_str)
                    for my_c in my_classes:
                        if my_c in c1 or c1 in my_c: return True
                    return False
                df_filtered = df_filtered[df_filtered['班级'].apply(class_match)]

            if df_filtered.empty:
                st.warning("⚠️ 在当前选择的群体或年级中，未找到您的授权班级数据。")
            else:
                is_single_subject_view = role in TEACHER_ROLES + SUBJECT_HEAD_ROLES
                
                # 常规一：首页
                if menu_sel == "首页":
                    total_stu = len(df_filtered)
                    class_count = df_filtered['班级'].nunique()
                    metric_col = subject if (is_single_subject_view and subject in df_filtered.columns) else '总分'
                    if not is_single_subject_view:
                        avail_metrics = ['总分'] + [s for s in ['语文','数学','英语','物理','化学','生物','历史','政治','地理'] if s in df_filtered.columns and df_filtered[s].sum() > 0]
                        st.markdown("<p style='font-size: 14px; font-weight: bold; color: #64748B; margin-bottom: -10px;'>⚡ 切换全局大盘分析指标</p>", unsafe_allow_html=True)
                        c_sel, _ = st.columns([1, 4])
                        metric_col = c_sel.selectbox("隐藏下拉", avail_metrics, label_visibility="collapsed")
                        st.markdown("<br>", unsafe_allow_html=True)
                    else:
                        if subject not in df_filtered.columns: st.warning(f"当前考试未配置您的学科【{subject}】的数据。"); st.stop()
                    
                    overall_avg = df_filtered[metric_col].mean().round(1) if total_stu > 0 else 0
                    class_avgs = df_filtered.groupby('班级')[metric_col].mean()
                    top_class = class_avgs.idxmax() if not class_avgs.empty else "无"
                    
                    k1, k2, k3, k4 = st.columns(4)
                    k1.metric("所辖班级数", f"{class_count} 个")
                    k2.metric("覆盖学生人数", f"{total_stu} 人")
                    k3.metric(f"{metric_col}群像均分", f"{overall_avg} 分")
                    k4.metric(f"{metric_col}领跑班级", top_class)
                    st.markdown("<br>", unsafe_allow_html=True)
                    
                    if not class_avgs.empty:
                        df_bar = class_avgs.reset_index().round(1)
                        fig_bar = px.bar(df_bar, x='班级', y=metric_col, text_auto=True, color='班级', color_discrete_sequence=px.colors.qualitative.Pastel, title=f"🏆 各班级【{metric_col}】均分横向对比阵列")
                        fig_bar.update_traces(textposition='outside', width=0.35, textfont_size=14, marker_line_width=0)
                        fig_bar.add_hline(y=overall_avg, line_dash="dot", line_color="#EF4444", line_width=2, annotation_text=f"🎯 群像均分红线: {overall_avg:.1f}", annotation_position="top left", annotation_font=dict(color="#EF4444", size=13, weight="bold"))
                        y_max = df_bar[metric_col].max() * 1.15 if not df_bar.empty else 100
                        fig_bar.update_layout(dragmode=False, showlegend=False, yaxis_range=[0, y_max], plot_bgcolor='rgba(248, 250, 252, 0.5)', paper_bgcolor='white', margin=dict(t=60, b=30, l=30, r=30), xaxis_title=None, yaxis_title=None, yaxis=dict(showgrid=True, gridcolor='#F1F5F9', gridwidth=1))
                        with st.container(border=True): st.plotly_chart(fig_bar, use_container_width=True, config=CHART_CONFIG)

                    # 常规单考上线率
                    st.markdown("<p style='font-size: 16px; font-weight: bold; color: #0F172A; margin-top: 30px;'>🎯 【目标考核】单次考试上线率达标监控</p>", unsafe_allow_html=True)
                    with st.container(border=True):
                        c_l1, c_l2, c_l3 = st.columns([1, 1, 2])
                        line_te = c_l1.number_input(f"🏆 设定【{metric_col}】特控线", min_value=0.0, value=0.0, step=1.0)
                        line_ben = c_l2.number_input(f"🎓 设定【{metric_col}】本科线", min_value=0.0, value=0.0, step=1.0)

                        if line_te > 0 or line_ben > 0:
                            pass_te_all = len(df_filtered[df_filtered[metric_col] >= line_te]) if line_te > 0 else 0
                            pass_ben_all = len(df_filtered[df_filtered[metric_col] >= line_ben]) if line_ben > 0 else 0
                            rate_te_all = round(pass_te_all / total_stu * 100, 1) if total_stu > 0 else 0
                            rate_ben_all = round(pass_ben_all / total_stu * 100, 1) if total_stu > 0 else 0

                            st.markdown(f"""
                            <div style='background-color: #F8FAFC; padding: 15px; border-radius: 8px; border-left: 4px solid #10B981; margin-top: 5px; margin-bottom: 20px;'>
                                <span style='font-size: 15px; color: #334155;'><b>🌐 群体整体达标概况：</b></span><br>
                                <span style='font-size: 14px; color: #64748B;'>
                                🏆 特控上线总计 <b>{pass_te_all}</b> 人 (上线率 <b style='color:#10B981;'>{rate_te_all}%</b>) &nbsp;&nbsp;|&nbsp;&nbsp; 
                                🎓 本科上线总计 <b>{pass_ben_all}</b> 人 (上线率 <b style='color:#3B82F6;'>{rate_ben_all}%</b>)
                                </span>
                            </div>
                            """, unsafe_allow_html=True)

                            class_stats = []
                            for cls in sorted(df_filtered['班级'].unique()):
                                cdf = df_filtered[df_filtered['班级'] == cls]
                                c_total = len(cdf)
                                row_dict = {'班级': cls, '参考人数': c_total}
                                if line_te > 0:
                                    c_te = len(cdf[cdf[metric_col] >= line_te])
                                    row_dict['特控上线数'] = c_te
                                    row_dict['特控上线率(%)'] = round(c_te / c_total * 100, 1) if c_total > 0 else 0
                                if line_ben > 0:
                                    c_ben = len(cdf[cdf[metric_col] >= line_ben])
                                    row_dict['本科上线数'] = c_ben
                                    row_dict['本科上线率(%)'] = round(c_ben / c_total * 100, 1) if c_total > 0 else 0
                                class_stats.append(row_dict)

                            df_stats = pd.DataFrame(class_stats)
                            if line_te > 0: df_stats = df_stats.sort_values('特控上线率(%)', ascending=False)
                            else: df_stats = df_stats.sort_values('本科上线率(%)', ascending=False)

                            html_stats = "<div style='width: 100%; overflow-x: auto; border-radius: 8px; border: 1px solid #E8E8E8;'><table style='width: 100%; border-collapse: collapse; font-size: 14px; text-align: center;'>"
                            html_stats += "<tr><th style='background-color: #FAFAFA; padding: 10px; border-bottom: 2px solid #E8E8E8; color: #333;'>班级</th><th style='background-color: #FAFAFA; padding: 10px; border-bottom: 2px solid #E8E8E8; color: #333;'>参考人数</th>"
                            if line_te > 0: html_stats += "<th style='background-color: #FFFBEB; padding: 10px; border-bottom: 2px solid #E8E8E8; color: #92400E;'>特控上线数</th><th style='background-color: #FFFBEB; padding: 10px; border-bottom: 2px solid #E8E8E8; color: #92400E;'>特控上线率</th>"
                            if line_ben > 0: html_stats += "<th style='background-color: #F0F9FF; padding: 10px; border-bottom: 2px solid #E8E8E8; color: #0369A1;'>本科上线数</th><th style='background-color: #F0F9FF; padding: 10px; border-bottom: 2px solid #E8E8E8; color: #0369A1;'>本科上线率</th>"
                            html_stats += "</tr>"

                            for _, r in df_stats.iterrows():
                                html_stats += f"<tr><td style='padding: 8px; border-bottom: 1px solid #F0F0F0;'><b>{r['班级']}</b></td><td style='padding: 8px; border-bottom: 1px solid #F0F0F0;'>{r['参考人数']}</td>"
                                if line_te > 0: html_stats += f"<td style='padding: 8px; border-bottom: 1px solid #F0F0F0; color:#B45309; font-weight:bold;'>{r['特控上线数']}</td><td style='padding: 8px; border-bottom: 1px solid #F0F0F0; color:#B45309;'>{r['特控上线率(%)']}%</td>"
                                if line_ben > 0: html_stats += f"<td style='padding: 8px; border-bottom: 1px solid #F0F0F0; color:#0284C7; font-weight:bold;'>{r['本科上线数']}</td><td style='padding: 8px; border-bottom: 1px solid #F0F0F0; color:#0284C7;'>{r['本科上线率(%)']}%</td>"
                                html_stats += "</tr>"
                            html_stats += "</table></div>"

                            st.markdown(html_stats, unsafe_allow_html=True)
                            file_data, file_name, mime_type = generate_excel_download(df_stats, f"上线率统计_{metric_col}", f"【{LATEST_EXAM['name']}】{metric_col}上线率报表")
                            st.download_button(label="📥 下载各班上线率统计报表 (Excel)", data=file_data, file_name=file_name, mime=mime_type, type="secondary")
                        else: st.info(f"👆 请在上方输入框内设定【{metric_col}】的「特控线」或「本科线」，系统将立即为您进行群体及各班的达标率动态测算。")

                # 常规二：成绩明细表
                elif menu_sel == "成绩明细表":
                    if is_single_subject_view:
                        st.info(f"💡 当前为学科专属视图，仅展示【{subject}】的单科成绩及排名。")
                        if subject in df_filtered.columns:
                            df_filtered[f'{subject}班级排名'] = df_filtered.groupby('班级')[subject].rank(ascending=False, method='min').fillna(0).astype(int)
                            df_filtered[f'{subject}年级排名'] = df_filtered[subject].rank(ascending=False, method='min').fillna(0).astype(int)
                            table_to_show = df_filtered[['姓名', '考号', '班级', f'{subject}年级排名', f'{subject}班级排名', subject]].sort_values(by=subject, ascending=False)
                            render_html_table(table_to_show)
                            file_data, file_name, mime_type = generate_excel_download(table_to_show, f"【{st.session_state.current_grade}】_{subject}明细", f"【{LATEST_EXAM['name']}】{subject}成绩单", split_by_class=True)
                            st.download_button(label="📥 下载精美 Excel 成绩单 (内含各班分表)", data=file_data, file_name=file_name, mime=mime_type, type="primary")
                        else: st.warning(f"当前考试中未找到您的学科【{subject}】。")
                    else:
                        st.info("💡 当前为全科汇总视图。")
                        cols = df_filtered.columns.tolist()
                        front_cols = ['姓名', '考号', '班级', '总分', '总分年级排名', '总分班级排名', '方向']
                        other_cols = [c for c in cols if c not in front_cols]
                        table_to_show = df_filtered[front_cols + other_cols].sort_values(by=['班级', '总分'], ascending=[True, False])
                        render_html_table(table_to_show)
                        file_data, file_name, mime_type = generate_excel_download(table_to_show, f"【{st.session_state.current_grade}】_{adm_direction}全科明细", f"【{LATEST_EXAM['name']}】{adm_direction}成绩汇总单", split_by_class=True)
                        st.download_button(label="📥 下载全科 Excel 汇总单 (内含各班分表)", data=file_data, file_name=file_name, mime=mime_type, type="primary")

                # 常规三：历次追踪分析
                elif menu_sel == "历次追踪分析":
                    st.info("🔍 在下方下拉框中搜索学生姓名，系统将跨越“时间胶囊”自动聚合该生的所有历史轨迹！")
                    student_options = df_filtered.apply(lambda x: f"{x['班级']} | {x['姓名']} | 考号:{x['考号']}", axis=1).tolist()
                    if student_options:
                        st.markdown("<p style='font-size: 14px; font-weight: bold; color: #64748B; margin-bottom: -10px;'>🔎 检索目标学生</p>", unsafe_allow_html=True)
                        c_sel, _ = st.columns([1, 1])
                        sel_student_str = c_sel.selectbox("隐藏标签", ["-- 请点击输入或选择学生 --"] + student_options, label_visibility="collapsed")
                        if sel_student_str != "-- 请点击输入或选择学生 --":
                            sel_id = clean_str(sel_student_str.split("考号:")[1].strip())
                            sel_name = sel_student_str.split("|")[1].strip()
                            history_records = []
                            with st.spinner(f"正在从云端数据库抽取【{sel_name}】的历史档案..."):
                                for i, exam in enumerate(EXAMS_LIST):
                                    dfs_sub = []
                                    subs_hist = ['语文','数学','英语','物理','化学','生物','历史','政治','地理']
                                    for sh in subs_hist:
                                        u = exam.get(sh)
                                        if u:
                                            d_sub = load_data(u, header_lines=[0,1,2])
                                            if d_sub is not None:
                                                n_c, i_c = None, None
                                                for col in d_sub.columns:
                                                    cstr = str(col[0]) if isinstance(col, tuple) else str(col)
                                                    if '姓名' in cstr: n_c = col
                                                    elif '考号' in cstr or '学号' in cstr: i_c = col
                                                
                                                if n_c:
                                                    r_ls = []
                                                    for _, row in d_sub.iterrows():
                                                        test_id = clean_str(row[i_c]) if i_c and pd.notna(row[i_c]) else clean_name(row[n_c])
                                                        if test_id == sel_id or test_id == sel_name:
                                                            tot = sum(float(row[c]) for c in d_sub.columns if c not in [n_c, i_c] and '总分' not in str(c) and '排名' not in str(c) and pd.notna(row[c]) and str(row[c]).replace('.','',1).isdigit())
                                                            r_ls.append({'考号': sel_id, sh: round(tot, 1)})
                                                    if r_ls: dfs_sub.append(pd.DataFrame(r_ls))
                                    if dfs_sub:
                                        exam_master = functools.reduce(lambda l, r: pd.merge(l, r, on=['考号'], how='outer'), dfs_sub)
                                        p_subs = [s for s in subs_hist if s in exam_master.columns]
                                        exam_master['总分'] = exam_master[p_subs].fillna(0).sum(axis=1).round(1)
                                        row = exam_master.iloc[0]
                                        rec = {"考试名称": exam['name'], "总分": float(row.get('总分', 0))}
                                        for s in subs_hist:
                                            if s in row and pd.notna(row[s]): rec[s] = float(row[s])
                                        history_records.append(rec)
                            if history_records:
                                st.markdown("<br>", unsafe_allow_html=True)
                                df_hist = pd.DataFrame(history_records)
                                with st.container(border=True):
                                    if is_single_subject_view:
                                        if subject in df_hist.columns:
                                            fig = px.line(df_hist, x="考试名称", y=subject, markers=True, title=f"📈 【{sel_name}】{subject} 历次成绩波动曲线", line_shape="spline")
                                            fig.update_traces(line_color="#3B82F6", marker=dict(size=10))
                                            fig.update_layout(dragmode=False, paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)', yaxis=dict(showgrid=True, gridcolor='#F1F5F9'))
                                            st.plotly_chart(fig, use_container_width=True, config=CHART_CONFIG)
                                        else: st.warning(f"未检索到该生【{subject}】的历史成绩。")
                                    else:
                                        t_col1, t_col2 = st.columns(2)
                                        with t_col1:
                                            if "总分" in df_hist.columns:
                                                fig1 = px.line(df_hist, x="考试名称", y="总分", markers=True, title=f"📈 【{sel_name}】历次总分波动曲线", line_shape="spline")
                                                fig1.update_traces(line_color="#3B82F6", marker=dict(size=10))
                                                fig1.update_layout(dragmode=False, paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)', yaxis=dict(showgrid=True, gridcolor='#F1F5F9'))
                                                st.plotly_chart(fig1, use_container_width=True, config=CHART_CONFIG)
                                        with t_col2:
                                            avail_hist_subs = [s for s in ['语文','数学','英语','物理','化学','生物','历史','政治','地理'] if s in df_hist.columns]
                                            if avail_hist_subs:
                                                st.markdown("<p style='font-size: 14px; font-weight: bold; color: #64748B; margin-bottom: -10px;'>🔬 透视单科走势</p>", unsafe_allow_html=True)
                                                sel_hist_sub = st.selectbox("隐藏标签", avail_hist_subs, key="hist_sub", label_visibility="collapsed")
                                                fig3 = px.line(df_hist, x="考试名称", y=sel_hist_sub, markers=True, title=f"📉 【{sel_name}】{sel_hist_sub} 单科走势", line_shape="spline")
                                                fig3.update_traces(line_color="#10B981", marker=dict(size=10))
                                                fig3.update_layout(dragmode=False, paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)', yaxis=dict(showgrid=True, gridcolor='#F1F5F9'))
                                                st.plotly_chart(fig3, use_container_width=True, config=CHART_CONFIG)

                # 常规四：AI 教研中心
                elif menu_sel == "AI 教研中心":
                    st.info("🧠 欢迎进入 AI 教研舱。系统将自动抓取底层题库，计算全班单题得分率，并进行智能聚类！")
                    analyze_subject = subject
                    if not is_single_subject_view:
                        avail_subs = [s for s in ['语文','数学','英语','物理','化学','生物','历史','政治','地理'] if s in df_filtered.columns]
                        if avail_subs:
                            st.markdown("<p style='font-size: 14px; font-weight: bold; color: #64748B; margin-bottom: -10px;'>⚙️ 选择要进行 AI 诊断的学科</p>", unsafe_allow_html=True)
                            c_sel, _ = st.columns([1, 3])
                            analyze_subject = c_sel.selectbox("隐藏下拉", avail_subs, label_visibility="collapsed")
                        else: analyze_subject = None
                    
                    if analyze_subject and LATEST_EXAM and LATEST_EXAM.get(analyze_subject):
                        df_diag = load_data(LATEST_EXAM[analyze_subject], header_lines=[0, 1, 2])
                        if df_diag is not None:
                            name_c, id_c, cls_c = None, None, None
                            for col in df_diag.columns:
                                cstr = str(col[0]) if isinstance(col, tuple) else str(col)
                                if '姓名' in cstr: name_c = col
                                elif '考号' in cstr or '学号' in cstr: id_c = col
                                elif '班级' in cstr: cls_c = col
                            if name_c:
                                def is_my_scope(c_val):
                                    c1 = normalize_class_name(c_val)
                                    if role in GLOBAL_ROLES + SUBJECT_HEAD_ROLES: return c1 in df_filtered['班级'].tolist()
                                    for my_c in my_classes:
                                        if my_c in c1 or c1 in my_c: return True
                                    return False
                                
                                # 修复：处理没有班级列的情况
                                if cls_c: df_diag_my = df_diag[df_diag[cls_c].apply(is_my_scope)]
                                else: df_diag_my = df_diag
                                
                                if not df_diag_my.empty:
                                    k_stats, weak_group_map = {}, {}
                                    for col in df_diag_my.columns:
                                        if col in [name_c, id_c, cls_c]: continue
                                        cstr = str(col[0]) if isinstance(col, tuple) else str(col)
                                        if '总分' in cstr or '排名' in cstr: continue
                                        q_name = str(col[0]).strip() if isinstance(col, tuple) else str(col).strip()
                                        kp = str(col[1]).strip() if isinstance(col, tuple) and len(col) > 1 else q_name
                                        if kp == "" or kp.startswith("Unnamed"): kp = q_name
                                        try: full = float(col[2]) if isinstance(col, tuple) and len(col) > 2 else 0
                                        except: full = 0
                                        if full <= 0:
                                            try: full = float(pd.to_numeric(df_diag_my[col], errors='coerce').max())
                                            except: full = 0
                                        if full > 0:
                                            if kp not in k_stats: k_stats[kp] = []
                                            k_stats[kp].append(pd.to_numeric(df_diag_my[col], errors='coerce').mean() / full)
                                            for _, r_data in df_diag_my.iterrows():
                                                stu_n = clean_name(r_data[name_c])
                                                try: score = float(r_data[col])
                                                except: score = 0
                                                if score < full * 0.6: 
                                                    if kp not in weak_group_map: weak_group_map[kp] = []
                                                    if stu_n and stu_n not in weak_group_map[kp]: weak_group_map[kp].append(stu_n)
                                    if k_stats:
                                        with st.container(border=True):
                                            k_final = [{"知识点": kp, "群体整体掌握率": round(sum(rates)/len(rates)*100, 1)} for kp, rates in k_stats.items()]
                                            df_k = pd.DataFrame(k_final).sort_values("群体整体掌握率")
                                            fig_k = px.bar(df_k, x="群体整体掌握率", y="知识点", orientation='h', title=f"🎯 【{analyze_subject}】知识点群体掌握率扫描")
                                            fig_k.update_layout(dragmode=False, paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)')
                                            st.plotly_chart(fig_k, use_container_width=True, config=CHART_CONFIG)
                                        
                                        sorted_weak_kps = sorted(weak_group_map.items(), key=lambda x: len(x[1]), reverse=True)
                                        grouped_str_list = []
                                        for kp, students in sorted_weak_kps[:5]:
                                            if students: grouped_str_list.append(f"【{kp}】得分率<60%的薄弱名单（需重点关注）：{', '.join(students)}")
                                        grouped_data_str = "\n".join(grouped_str_list)
                                        
                                        ai_t_key = f"ai_tea_{st.session_state.current_grade}_{name}_{analyze_subject}"
                                        if AI_API_KEY:
                                            if st.button("✨ 一键生成【AI 分层教学与靶向辅导报告】", type="primary"):
                                                with st.spinner("AI 大脑正在深度剖析每个学生的单题得分，进行聚类提取..."):
                                                    ai_reply = get_ai_grouped_advice_for_teacher(st.session_state.current_grade, analyze_subject, grouped_data_str)
                                                    st.session_state[ai_t_key] = ai_reply

                                            if ai_t_key in st.session_state:
                                                saved_reply = st.session_state[ai_t_key]
                                                st.markdown(f"<div class='ai-box'><b>🤖 专家指导建议：</b><br><br>{saved_reply}</div><br>", unsafe_allow_html=True)
                                                t_c1, t_c2, t_c3 = st.columns([1.5, 1, 1])
                                                with t_c1:
                                                    if st.button("📌 存入本设备暂存库"):
                                                        saved_t_list_key = f"saved_ai_tea_list_{st.session_state.current_grade}_{analyze_subject}"
                                                        if saved_t_list_key not in st.session_state: st.session_state[saved_t_list_key] = []
                                                        st.session_state[saved_t_list_key].insert(0, saved_reply)
                                                        st.toast("✅ 已成功存入！")
                                                with t_c3:
                                                    doc_title = f"【{LATEST_EXAM['name']}】{analyze_subject}_AI教研报告"
                                                    file_data, file_name, mime_type = generate_ai_doc(doc_title, saved_reply)
                                                    st.download_button(label="📥 下载 Word 排版报告", data=file_data, file_name=file_name, mime=mime_type, type="primary")
                    else: st.warning(f"目前缺少【{analyze_subject}】的单题明细表，无法进行 AI 深度诊断。")
